import os
import tkinter as tk
from tkinter import filedialog, messagebox, simpledialog, ttk
from docx import Document
import PyPDF2
import pandas as pd
import subprocess
import platform
from pdfminer.high_level import extract_text
import sqlite3
import time
import threading
import json
import re
import numpy as np
# Add dotenv support
try:
    from dotenv import load_dotenv
    load_dotenv()
except ImportError:
    pass

# AI Search imports
try:
    from ai_search import initialize_ai_search, add_documents_to_ai_index, ai_search, clear_ai_index, get_ai_index_stats
    AI_AVAILABLE = True
    print("Full AI search available!")
except ImportError:
    AI_AVAILABLE = False
    print("AI search dependencies not available. Install with: pip install sentence-transformers chromadb torch")

INDEX_DB = os.path.join(os.path.dirname(__file__), 'file_index.db')
search_cancelled = False
indexing_in_progress = False

# === Helper for loading embeddings ===
def load_embeddings(filename):
    try:
        with open(filename, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception as e:
        print(f"[Failed to load embeddings: {e}]")
        return {}

# === Database Schema Migration ===
def ensure_schema():
    conn = sqlite3.connect(INDEX_DB, timeout=30)
    c = conn.cursor()
    # Check if file_index exists and has root_path column
    c.execute("SELECT name FROM sqlite_master WHERE type='table' AND name='file_index'")
    if c.fetchone():
        try:
            c.execute('SELECT root_path FROM file_index LIMIT 1')
        except sqlite3.OperationalError:
            c.execute('DROP TABLE file_index')
    c.execute('''CREATE TABLE IF NOT EXISTS roots (root_path TEXT PRIMARY KEY)''')
    c.execute('''CREATE VIRTUAL TABLE IF NOT EXISTS file_index USING fts5(
        file_path, file_type, mtime, content, root_path
    )''')
    conn.commit()
    conn.close()

# === Multiple Roots Management ===
def init_db():
    ensure_schema()

def get_roots():
    init_db()
    conn = sqlite3.connect(INDEX_DB, timeout=30)
    c = conn.cursor()
    c.execute('SELECT root_path FROM roots')
    roots = [row[0] for row in c.fetchall()]
    conn.close()
    return roots

def add_root(root_path):
    init_db()
    conn = sqlite3.connect(INDEX_DB, timeout=30)
    c = conn.cursor()
    c.execute('INSERT OR IGNORE INTO roots (root_path) VALUES (?)', (root_path,))
    conn.commit()
    conn.close()

def remove_root(root_path):
    init_db()
    conn = sqlite3.connect(INDEX_DB, timeout=30)
    c = conn.cursor()
    c.execute('DELETE FROM roots WHERE root_path=?', (root_path,))
    c.execute('DELETE FROM file_index WHERE root_path=?', (root_path,))
    conn.commit()
    conn.close()

# === Indexing Functions (now for all roots) ===
def extract_file_content(file_path):
    try:
        if file_path.endswith('.txt'):
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
        elif file_path.endswith('.md'):
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
        elif file_path.endswith('.docx'):
            doc = Document(file_path)
            return '\n'.join([para.text for para in doc.paragraphs])
        elif file_path.endswith('.pdf'):
            return extract_text(file_path)
        elif file_path.endswith('.xlsx') and not os.path.basename(file_path).startswith('~$'):
            df = pd.read_excel(file_path, sheet_name=None)
            content = []
            for sheet_name, sheet_data in df.items():
                for row in sheet_data.itertuples(index=False):
                    content.append(' '.join([str(cell) for cell in row if pd.notnull(cell)]))
            return '\n'.join(content)
    except Exception as e:
        print(f"Error extracting {file_path}: {e}")
    return ''

def get_selected_roots():
    selected = roots_listbox.curselection()
    if selected:
        return [roots_listbox.get(i) for i in selected]
    else:
        return get_roots()

def build_index_all():
    global search_cancelled, indexing_in_progress
    if indexing_in_progress:
        return  # Skip if another operation is in progress
    indexing_in_progress = True
    search_cancelled = False
    try:
        init_db()
        roots = get_selected_roots()
        conn = sqlite3.connect(INDEX_DB, timeout=30)
        c = conn.cursor()
        c.execute('DELETE FROM file_index')
        conn.commit()
        for root_path in roots:
            for root, dirs, files in os.walk(root_path):
                for file in files:
                    if search_cancelled:
                        conn.commit()
                        conn.close()
                        indexing_in_progress = False
                        return
                    file_path = os.path.join(root, file)
                    if file.endswith(('.txt', '.md', '.docx', '.pdf', '.xlsx')) and not file.startswith('~$'):
                        mtime = os.path.getmtime(file_path)
                        content = extract_file_content(file_path)
                        c.execute('INSERT INTO file_index (file_path, file_type, mtime, content, root_path) VALUES (?, ?, ?, ?, ?)',
                                  (file_path, os.path.splitext(file)[1][1:].upper(), mtime, content, root_path))
        conn.commit()
        conn.close()
    finally:
        indexing_in_progress = False

def update_index_all():
    global search_cancelled, indexing_in_progress
    if indexing_in_progress:
        return  # Skip if another operation is in progress
    indexing_in_progress = True
    search_cancelled = False
    try:
        init_db()
        roots = get_selected_roots()
        conn = sqlite3.connect(INDEX_DB, timeout=30)
        c = conn.cursor()
        c.execute('SELECT file_path, mtime, root_path FROM file_index')
        indexed = {(row[0], row[2]): row[1] for row in c.fetchall()}
        seen = set()
        for root_path in roots:
            for root, dirs, files in os.walk(root_path):
                for file in files:
                    if search_cancelled:
                        conn.commit()
                        conn.close()
                        indexing_in_progress = False
                        return
                    file_path = os.path.join(root, file)
                    if file.endswith(('.txt', '.md', '.docx', '.pdf', '.xlsx')) and not file.startswith('~$'):
                        mtime = os.path.getmtime(file_path)
                        seen.add((file_path, root_path))
                        if (file_path, root_path) not in indexed:
                            content = extract_file_content(file_path)
                            c.execute('INSERT INTO file_index (file_path, file_type, mtime, content, root_path) VALUES (?, ?, ?, ?, ?)',
                                      (file_path, os.path.splitext(file)[1][1:].upper(), mtime, content, root_path))
                        elif float(indexed[(file_path, root_path)]) < mtime:
                            content = extract_file_content(file_path)
                            c.execute('UPDATE file_index SET mtime=?, content=? WHERE file_path=? AND root_path=?',
                                      (mtime, content, file_path, root_path))
        for (file_path, root_path) in indexed:
            if (file_path, root_path) not in seen:
                c.execute('DELETE FROM file_index WHERE file_path=? AND root_path=?', (file_path, root_path))
        conn.commit()
        conn.close()
    finally:
        indexing_in_progress = False

def search_index(keyword):
    # Get all roots and selected roots BEFORE opening any DB connection
    all_roots = get_roots()
    selected_roots = get_selected_roots()
    results = []
    # Now open the DB connection
    conn = sqlite3.connect(INDEX_DB, timeout=30)
    c = conn.cursor()
    if selected_roots and len(selected_roots) < len(all_roots):
        placeholders = ','.join('?' for _ in selected_roots)
        q = f"SELECT file_path, file_type, snippet(file_index, 3, '[', ']', '...', 20), root_path FROM file_index WHERE content MATCH ? AND root_path IN ({placeholders})"
        c.execute(q, (keyword, *selected_roots))
    else:
        q = f"SELECT file_path, file_type, snippet(file_index, 3, '[', ']', '...', 20), root_path FROM file_index WHERE content MATCH ?"
        c.execute(q, (keyword,))
    for file_path, file_type, snippet_, root_path in c.fetchall():
        results.append({
            "File Path": file_path,
            "File Type": file_type,
            "Location": f"Indexed ({root_path})",
            "Content": snippet_
        })
    conn.close()
    return results

# === Live Search Functions ===
def search_txt(file_path, keyword, results):
    global search_cancelled
    try:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            for i, line in enumerate(f, start=1):
                if search_cancelled:
                    return
                if keyword.lower() in line.lower():
                    results.append({
                        "File Path": file_path,
                        "File Type": "TXT",
                        "Location": f"Line {i}",
                        "Content": line.strip()
                    })
    except Exception as e:
        print(f"Error reading TXT {file_path}: {e}")

def search_md(file_path, keyword, results):
    global search_cancelled
    try:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            for i, line in enumerate(f, start=1):
                if search_cancelled:
                    return
                if keyword.lower() in line.lower():
                    results.append({
                        "File Path": file_path,
                        "File Type": "MD",
                        "Location": f"Line {i}",
                        "Content": line.strip()
                    })
    except Exception as e:
        print(f"Error reading MD {file_path}: {e}")

def search_docx(file_path, keyword, results):
    global search_cancelled
    try:
        doc = Document(file_path)
        for i, para in enumerate(doc.paragraphs, start=1):
            if search_cancelled:
                return
            if keyword.lower() in para.text.lower():
                results.append({
                    "File Path": file_path,
                    "File Type": "DOCX",
                    "Location": f"Paragraph {i}",
                    "Content": para.text.strip()
                })
    except Exception as e:
        print(f"Error reading DOCX {file_path}: {e}")

def search_pdf(file_path, keyword, results):
    global search_cancelled
    try:
        text = extract_text(file_path)
        if search_cancelled:
            return
        if text and keyword.lower() in text.lower():
            results.append({
                "File Path": file_path,
                "File Type": "PDF",
                "Location": "Full Text",
                "Content": "Match found"
            })
    except Exception as e:
        print(f"Error reading PDF {file_path}: {e}")

def search_xlsx(file_path, keyword, results):
    global search_cancelled
    try:
        df = pd.read_excel(file_path, sheet_name=None) # All sheets
        for sheet_name, sheet_data in df.items():
            for row_idx, row in sheet_data.iterrows():
                if search_cancelled:
                    return
                for col_name, cell_value in row.items():
                    if pd.notnull(cell_value) and keyword.lower() in str(cell_value).lower():
                        # Convert row_idx to string first, then to int to handle various index types
                        try:
                            row_num = int(str(row_idx)) + 1
                        except (ValueError, TypeError):
                            row_num = 1  # Fallback if conversion fails
                        results.append({
                            "File Path": file_path,
                            "File Type": "XLSX",
                            "Location": f"Sheet {sheet_name}, Row {row_num}, Column {col_name}",
                            "Content": str(cell_value)
                        })
    except Exception as e:
        print(f"Error reading XLSX {file_path}: {e}")

def search_folder(folder_path, keyword, results):
    global search_cancelled
    search_cancelled = False
    for root, dirs, files in os.walk(folder_path):
        for file in files:
            if search_cancelled:
                return
            file_path = os.path.join(root, file)
            if file.endswith('.txt'):
                search_txt(file_path, keyword, results)
            elif file.endswith('.md'):
                search_md(file_path, keyword, results)
            elif file.endswith('.docx'):
                search_docx(file_path, keyword, results)
            elif file.endswith('.pdf'):
                search_pdf(file_path, keyword, results)
            elif file.endswith('.xlsx'):
                search_xlsx(file_path, keyword, results)

# === GUI Functions for Roots ===
def add_root_gui():
    folder = filedialog.askdirectory()
    if folder:
        add_root(folder)
        update_roots_listbox()

def remove_root_gui():
    selection = roots_listbox.curselection()
    if selection:
        roots = [roots_listbox.get(i) for i in selection]
        if messagebox.askyesno("Remove Roots", f"Remove selected roots and all their files from index?\n{', '.join(roots)}"):
            wait_win = show_wait_message("Removing root(s), please wait...")
            for root in roots:
                remove_root(root)
            wait_win.destroy()
            update_roots_listbox()

def move_root_up():
    selection = roots_listbox.curselection()
    if not selection or selection[0] == 0:
        return  # Can't move up if nothing selected or already at top
    
    selected_index = selection[0]
    roots = get_roots()
    if selected_index > 0:
        # Swap the selected root with the one above it
        roots[selected_index], roots[selected_index - 1] = roots[selected_index - 1], roots[selected_index]
        
        # Update the database with new order
        conn = sqlite3.connect(INDEX_DB, timeout=30)
        c = conn.cursor()
        c.execute('DELETE FROM roots')  # Clear all roots
        for root in roots:
            c.execute('INSERT INTO roots (root_path) VALUES (?)', (root,))
        conn.commit()
        conn.close()
        
        update_roots_listbox()
        # Reselect the moved item
        roots_listbox.selection_set(selected_index - 1)

def move_root_down():
    selection = roots_listbox.curselection()
    if not selection:
        return  # Can't move down if nothing selected
    
    selected_index = selection[0]
    roots = get_roots()
    if selected_index < len(roots) - 1:
        # Swap the selected root with the one below it
        roots[selected_index], roots[selected_index + 1] = roots[selected_index + 1], roots[selected_index]
        
        # Update the database with new order
        conn = sqlite3.connect(INDEX_DB, timeout=30)
        c = conn.cursor()
        c.execute('DELETE FROM roots')  # Clear all roots
        for root in roots:
            c.execute('INSERT INTO roots (root_path) VALUES (?)', (root,))
        conn.commit()
        conn.close()
        
        update_roots_listbox()
        # Reselect the moved item
        roots_listbox.selection_set(selected_index + 1)

def update_roots_listbox():
    roots_listbox.delete(0, tk.END)
    for root in get_roots():
        roots_listbox.insert(tk.END, root)

# === Threaded Indexing/Search Functions ===
def build_index_all_thread():
    thread = threading.Thread(target=build_index_all_gui)
    thread.start()

def show_wait_message(msg):
    wait_win = tk.Toplevel(root)
    wait_win.title("Please wait")
    wait_win.geometry("300x80")
    wait_win.resizable(False, False)
    wait_win.grab_set()
    tk.Label(wait_win, text=msg, font=("Arial", 12)).pack(expand=True, fill="both", padx=10, pady=20)
    wait_win.update()
    return wait_win

def build_index_all_gui():
    global search_cancelled
    search_cancelled = False
    wait_win = show_wait_message("Rebuilding index, please wait...")
    t0 = time.time()
    build_index_all()
    wait_win.destroy()
    root.after(0, lambda: messagebox.showinfo("Index", f"Index rebuilt in {time.time() - t0:.1f} seconds."))

def update_index_all_thread():
    thread = threading.Thread(target=update_index_all_gui)
    thread.start()

def update_index_all_gui():
    global search_cancelled
    search_cancelled = False
    wait_win = show_wait_message("Updating index, please wait...")
    t0 = time.time()
    update_index_all()
    wait_win.destroy()
    root.after(0, lambda: messagebox.showinfo("Index", f"Index updated in {time.time() - t0:.1f} seconds."))

def update_index_periodically():
    global search_cancelled
    search_cancelled = False
    threading.Thread(target=update_index_all).start()
    root.after(30 * 60 * 1000, update_index_periodically)  # every 30 minutes

def start_live_search_thread():
    thread = threading.Thread(target=start_live_search)
    thread.start()

def start_live_search():
    global search_cancelled
    search_cancelled = False
    keyword = keyword_entry.get()
    if not keyword:
        root.after(0, lambda: messagebox.showwarning("Input Error", "Please enter a keyword."))
        return
    results.clear()
    for root_path in get_selected_roots():
        search_folder(root_path, keyword, results)
    root.after(0, lambda: show_results(results))

def show_results(results):
    for widget in results_inner_frame.winfo_children():
        widget.destroy()
    
    # Add bundle toggle checkbox
    bundle_frame = tk.Frame(results_inner_frame, bg="#f0f0f0")
    bundle_frame.pack(fill="x", padx=5, pady=(5,0))
    tk.Checkbutton(bundle_frame, text="Bundle by file (show only best match per file)", variable=bundle_by_file, bg="#f0f0f0", command=lambda: show_results(results)).pack(side="left", padx=5, pady=5)

    display_results = results
    if bundle_by_file.get():
        # Deduplicate: keep only the best match per file (first occurrence or highest score if available)
        file_best = {}
        for res in results:
            file_path = res['File Path']
            # Use similarity score if available, else just first occurrence
            score = None
            if 'Location' in res and 'Score:' in res['Location']:
                try:
                    score = float(res['Location'].split('Score:')[1].split(')')[0])
                except:
                    score = None
            if file_path not in file_best:
                file_best[file_path] = (res, score)
            else:
                prev_res, prev_score = file_best[file_path]
                if score is not None and (prev_score is None or score > prev_score):
                    file_best[file_path] = (res, score)
        display_results = [v[0] for v in file_best.values()]

    if display_results:
        # Create header row
        header_frame = tk.Frame(results_inner_frame, bg="#f0f0f0", relief="raised", bd=1)
        header_frame.pack(fill="x", padx=5, pady=(5,0))
        
        tk.Label(header_frame, text="File Type", font=("Arial", 10, "bold"), bg="#f0f0f0", width=8).pack(side="left", padx=5, pady=5)
        tk.Label(header_frame, text="File Path", font=("Arial", 10, "bold"), bg="#f0f0f0", width=30, anchor="w").pack(side="left", padx=5, pady=5, fill="x", expand=True)
        tk.Label(header_frame, text="Location", font=("Arial", 10, "bold"), bg="#f0f0f0", width=15).pack(side="left", padx=5, pady=5)
        tk.Label(header_frame, text="Content", font=("Arial", 10, "bold"), bg="#f0f0f0", width=30, anchor="w").pack(side="left", padx=5, pady=5)
        tk.Label(header_frame, text="Actions", font=("Arial", 10, "bold"), bg="#f0f0f0", width=15).pack(side="left", padx=5, pady=5)
        
        # Tooltip for file path
        tooltip = tk.Toplevel(root)
        tooltip.withdraw()
        tooltip.overrideredirect(True)
        tooltip_label = tk.Label(tooltip, text="", background="#ffffe0", relief="solid", borderwidth=1, font=("Arial", 9))
        tooltip_label.pack(ipadx=1)
        
        def show_tooltip(event, text):
            tooltip_label.config(text=text)
            tooltip.deiconify()
            tooltip.lift()
            x = event.widget.winfo_rootx() + event.x + 20
            y = event.widget.winfo_rooty() + event.y + 10
            tooltip.geometry(f"+{x}+{y}")
        def hide_tooltip(event):
            tooltip.withdraw()
        
        # Display results with alternating row colors
        for idx, res in enumerate(display_results):
            row_bg = "#ffffff" if idx % 2 == 0 else "#f8f8f8"
            result_frame = tk.Frame(results_inner_frame, bg=row_bg, relief="flat", bd=1)
            result_frame.pack(fill="x", padx=5, pady=1)
            
            file_type_colors = {
                "TXT": "#4CAF50",
                "DOCX": "#2196F3",
                "PDF": "#FF5722",
                "XLSX": "#FF9800",
                "MD": "#9C27B0"
            }
            type_color = file_type_colors.get(res['File Type'], "#666666")
            tk.Label(result_frame, text=res['File Type'], font=("Arial", 9, "bold"), 
                    fg="white", bg=type_color, width=8, relief="raised").pack(side="left", padx=5, pady=5)
            
            path_text = res['File Path']
            display_path = os.path.basename(path_text)
            path_label = tk.Label(result_frame, text=display_path, font=("Arial", 9), 
                                  bg=row_bg, anchor="w", width=30)
            path_label.pack(side="left", padx=5, pady=5, fill="x", expand=True)
            path_label.bind("<Enter>", lambda e, t=path_text: show_tooltip(e, t))
            path_label.bind("<Leave>", hide_tooltip)
            
            # Location column improvement
            location_text = res['Location']
            # For AI results, show chunk info if available
            if 'chunk_index' in res and 'total_chunks' in res:
                location_text = f"Chunk {res['chunk_index']+1}/{res['total_chunks']} (Score: {res.get('similarity_score', 0):.2f})"
            elif 'Score:' in location_text and 'chunk' in res:
                location_text = f"Chunk {res['chunk']} (Score: {res.get('similarity_score', 0):.2f})"
            # For index search, try to extract more context if possible (future: line/paragraph)
            tk.Label(result_frame, text=location_text, font=("Arial", 9), 
                    bg=row_bg, anchor="w", width=15).pack(side="left", padx=5, pady=5)
            
            # Display content if available, with keyword highlighting
            content_text = res.get('Content', '')
            if content_text:
                # Always show enough content for highlighting
                display_content = content_text[:300] + "..." if len(content_text) > 300 else content_text
                # Remove all brackets for display
                clean_content = re.sub(r'\[([^\]]+)\]', r'\1', display_content)
                print(f"[DEBUG] Content for highlighting: {clean_content}")
                content_widget = tk.Text(result_frame, height=3, width=50, font=("Arial", 9), bg=row_bg, wrap="word", borderwidth=2, relief="solid", highlightthickness=2, highlightbackground="#888")
                content_widget.insert("1.0", clean_content)
                # Robustly highlight all words in the keyword/phrase
                keyword = keyword_entry.get().strip()
                print(f"[DEBUG] Keyword entry: '{keyword}'")
                match_count = 0
                if keyword:
                    words = [w for w in re.split(r'\s+', keyword) if w]
                    print(f"[DEBUG] Words to highlight: {words}")
                    for word in words:
                        for match in re.finditer(re.escape(word), clean_content, re.IGNORECASE):
                            start, end = match.start(), match.end()
                            start_idx = f"1.0+{start}c"
                            end_idx = f"1.0+{end}c"
                            print(f"[DEBUG] Highlighting '{clean_content[start:end]}' at {start_idx} to {end_idx}")
                            content_widget.tag_add("highlight", start_idx, end_idx)
                            match_count += 1
                    content_widget.tag_config("highlight", foreground="#228B22", font=("Arial", 9, "bold"))
                if match_count == 0:
                    print(f"[DEBUG] No matches found for highlighting in: {clean_content}")
                content_widget.config(state="disabled")
                content_widget.pack(side="left", padx=5, pady=5)
                content_widget.bind("<Enter>", lambda e, t=content_text: show_tooltip(e, t))
                content_widget.bind("<Leave>", hide_tooltip)
            else:
                tk.Label(result_frame, text="", font=("Arial", 9), 
                        bg=row_bg, anchor="w", width=30).pack(side="left", padx=5, pady=5)
            
            actions_frame = tk.Frame(result_frame, bg=row_bg)
            actions_frame.pack(side="left", padx=5, pady=5)
            tk.Button(actions_frame, text="📄 Open", command=lambda p=res['File Path']: open_file(p), 
                     bg="#4CAF50", fg="white", font=("Arial", 8, "bold"), 
                     relief="flat", bd=0, padx=8, pady=2).pack(side="left", padx=2)
            tk.Button(actions_frame, text="📁 Folder", command=lambda p=res['File Path']: open_folder_location(p), 
                     bg="#2196F3", fg="white", font=("Arial", 8, "bold"), 
                     relief="flat", bd=0, padx=8, pady=2).pack(side="left", padx=2)
        
        # Consolidated AI summary (with debug and fallback)
        is_ai_results = any('AI Match' in res.get('Location', '') or 'Score:' in res.get('Location', '') for res in results)
        if is_ai_results and AI_AVAILABLE and len(results) > 1:
            all_contents = []
            for res in results:
                content = res.get('Content', '')
                if content.startswith('📝 Summary:'):
                    content = content.split('\n\n📄 Content:')[-1]
                all_contents.append(content)
            all_text = '\n'.join(all_contents)
            summary_text = None
            try:
                model_choice = ai_model_var.get()
                summary_text = ai_summarize_dispatch(all_text, model_choice)
            except Exception as e:
                print(f"[DEBUG] AI summary generation failed: {e}")
            if not summary_text:
                # Fallback: concatenate top 3 results
                summary_text = '\n'.join(all_contents[:3])
            summary_frame = tk.Frame(results_inner_frame, bg="#e3f2fd", relief="groove", bd=2)
            summary_frame.pack(fill="x", padx=5, pady=(5,0))
            tk.Label(summary_frame, text="🤖 Consolidated AI Summary:", font=("Arial", 10, "bold"), fg="#1565C0", bg="#e3f2fd").pack(anchor="w", padx=10, pady=(5,0))
            tk.Label(summary_frame, text=summary_text, font=("Arial", 9), fg="#1565C0", bg="#e3f2fd", wraplength=800, justify="left").pack(anchor="w", padx=10, pady=(0,5))

        summary_frame = tk.Frame(results_inner_frame, bg="#e8f5e8", relief="groove", bd=1)
        summary_frame.pack(fill="x", padx=5, pady=(5,0))
        tk.Label(summary_frame, text=f"✓ Found {len(display_results)} match{'es' if len(display_results) != 1 else ''}", 
                font=("Arial", 10, "bold"), fg="#2E7D32", bg="#e8f5e8").pack(pady=5)
        messagebox.showinfo("Search Complete", f"Found {len(display_results)} matches.")
    else:
        no_results_frame = tk.Frame(results_inner_frame, bg="#fff3cd", relief="groove", bd=2)
        no_results_frame.pack(fill="x", padx=20, pady=20)
        tk.Label(no_results_frame, text="🔍 No matches found", 
                font=("Arial", 12, "bold"), fg="#856404", bg="#fff3cd").pack(pady=10)
        tk.Label(no_results_frame, text="Try different keywords or check your file types", 
                font=("Arial", 9), fg="#856404", bg="#fff3cd").pack(pady=(0,10))
        messagebox.showinfo("Search Complete", "No matches found.")

# Helper functions for context menu actions

def open_selected_file(tree):
    selection = tree.selection()
    if selection:
        item = tree.item(selection[0])
        file_path = item['values'][1]
        open_file(file_path)

def open_selected_folder(tree):
    selection = tree.selection()
    if selection:
        item = tree.item(selection[0])
        file_path = item['values'][1]
        open_folder_location(file_path)

def cancel_search():
    global search_cancelled
    search_cancelled = True
    root.after(0, lambda: messagebox.showinfo("Cancelled", "Operation cancelled by user."))

def start_index_search():
    global search_cancelled
    search_cancelled = False
    keyword = keyword_entry.get()
    if not keyword:
        root.after(0, lambda: messagebox.showwarning("Input Error", "Please enter a keyword."))
        return
    results.clear()
    results.extend(search_index(keyword))
    root.after(0, lambda: show_results(results))

def start_ai_search():
    global search_cancelled
    search_cancelled = False
    keyword = keyword_entry.get()
    if not keyword:
        root.after(0, lambda: messagebox.showwarning("Input Error", "Please enter a keyword."))
        return
    
    if not AI_AVAILABLE:
        root.after(0, lambda: messagebox.showerror("AI Search Error", "AI search is not available. Please install dependencies:\npip install sentence-transformers chromadb torch"))
        return
    
    results.clear()
    model_choice = ai_model_var.get()
    ai_results = ai_search_dispatch(keyword, n_results=20, model_choice=model_choice)
    
    # Filter by selected roots
    selected_roots = get_selected_roots()
    def is_in_selected_roots(path):
        import os
        return any(os.path.abspath(path).startswith(os.path.abspath(root)) for root in selected_roots)
    ai_results = [r for r in ai_results if is_in_selected_roots(r.get('file_path', ''))]
    
    # Convert AI results to standard format
    for result in ai_results:
        # Create enhanced content with summary if available
        content = result.get('content', '')
        content = content[:200] + "..." if len(content) > 200 else content
        if result.get('summary'):
            content = f"📝 Summary: {result['summary']}\n\n📄 Content: {content}"
        
        results.append({
            "File Path": result.get('file_path', ''),
            "File Type": result.get('file_type', ''),
            "Location": f"🤖 AI Match (Score: {result.get('similarity_score', 0):.2f})",
            "Content": content
        })
    
    root.after(0, lambda: show_results(results))

def build_ai_index():
    """Build AI search index from current indexed files"""
    if not AI_AVAILABLE:
        messagebox.showerror("AI Search Error", "AI search is not available. Please install dependencies:\npip install sentence-transformers chromadb torch")
        return
    
    # Get all indexed files
    conn = sqlite3.connect(INDEX_DB, timeout=30)
    c = conn.cursor()
    c.execute('SELECT file_path, file_type, content FROM file_index')
    files = c.fetchall()
    conn.close()
    
    if not files:
        messagebox.showinfo("AI Index", "No files found in index. Please build the regular index first.")
        return
    
    # Prepare documents for AI index
    documents = []
    for file_path, file_type, content in files:
        documents.append({
            'file_path': file_path,
            'file_type': file_type,
            'content': content
        })
    
    # Show progress dialog
    progress_win = show_wait_message("Building AI index, please wait...")
    
    def build_ai_index_thread():
        try:
            # Initialize AI search
            if not initialize_ai_search():
                try:
                    root.after(0, lambda: progress_win.destroy())
                    root.after(0, lambda: messagebox.showerror("AI Search Error", "Failed to initialize AI search engine"))
                except:
                    pass
                return
            
            # Add documents to AI index
            success = add_documents_to_ai_index(documents)
            
            try:
                root.after(0, lambda: progress_win.destroy())
                if success:
                    stats = get_ai_index_stats()
                    root.after(0, lambda: messagebox.showinfo("AI Index", f"AI index built successfully!\n\nDocuments: {stats['total_documents']}\nSize: {stats['index_size']}"))
                else:
                    root.after(0, lambda: messagebox.showerror("AI Index Error", "Failed to build AI index"))
            except:
                pass
                
        except Exception as e:
            try:
                root.after(0, lambda: progress_win.destroy())
                root.after(0, lambda: messagebox.showerror("AI Index Error", f"Error building AI index: {e}"))
            except:
                pass
    
    thread = threading.Thread(target=build_ai_index_thread)
    thread.start()

def clear_ai_index_gui():
    """Clear AI search index"""
    if not AI_AVAILABLE:
        messagebox.showerror("AI Search Error", "AI search is not available.")
        return
    
    if messagebox.askyesno("Clear AI Index", "Are you sure you want to clear the AI search index?"):
        success = clear_ai_index()
        if success:
            messagebox.showinfo("AI Index", "AI search index cleared successfully!")
        else:
            messagebox.showerror("AI Index Error", "Failed to clear AI index")

# === Open File Functions ===
def open_file(file_path):
    try:
        if not os.path.exists(file_path):
            messagebox.showerror("Error", f"File does not exist:\n{file_path}")
            return
        if platform.system() == "Windows":
            os.startfile(file_path)
        elif platform.system() == "Darwin":  # macOS
            subprocess.run(["open", file_path], check=True)
        else:  # Linux
            subprocess.run(["xdg-open", file_path], check=True)
    except Exception as e:
        messagebox.showerror("Error", f"Cannot open file:\n{file_path}\n\nError: {e}")

def open_folder_location(file_path):
    folder = os.path.dirname(file_path)
    try:
        if platform.system() == "Windows":
            os.startfile(folder)
        elif platform.system() == "Darwin":  # macOS
            subprocess.run(["open", "-R", file_path])
        else:  # Linux
            subprocess.run(["xdg-open", folder])
    except Exception as e:
        messagebox.showerror("Error", f"Cannot open folder: {e}")

def clear_results():
    for widget in results_inner_frame.winfo_children():
        widget.destroy()

# === App Window Layout ===
results = []

root = tk.Tk()
root.title("🔍 SearchAuto - Universal File Content Search")
root.configure(bg="#f5f5f5")  # Light gray background
root.minsize(1200, 700)  # Ensure both panels are visible
root.geometry("1400x800")  # Force initial window width for both panels

bundle_by_file = tk.BooleanVar(value=True)

# Roots Management Frame (top, full width)
roots_frame = tk.LabelFrame(root, text="📁 Indexed Roots", font=("Arial", 11, "bold"), fg="navy", relief="groove", bd=2)
roots_frame.grid(row=0, column=0, columnspan=2, padx=10, pady=8, sticky="ew")
tk.Label(roots_frame, text="Select roots to search (Ctrl+Click for multi-select):", font=("Arial", 9)).pack(anchor="w", padx=5, pady=2)

# Create a frame for the listbox and arrow buttons
roots_content_frame = tk.Frame(roots_frame)
roots_content_frame.pack(fill="x", expand=True, padx=5, pady=5)

# Listbox on the left
roots_listbox = tk.Listbox(roots_content_frame, width=80, height=3, selectmode=tk.MULTIPLE, font=("Arial", 9), relief="flat", bd=1)
roots_listbox.pack(side="left", fill="x", expand=True, padx=(0,5))

# Arrow buttons frame in the middle
roots_arrow_frame = tk.Frame(roots_content_frame)
roots_arrow_frame.pack(side="left", fill="y", padx=5)
tk.Button(roots_arrow_frame, text="▲", command=move_root_up, bg="#2196F3", fg="white", font=("Arial", 12, "bold"), 
          relief="flat", bd=0, width=3, height=1).pack(pady=2)
tk.Button(roots_arrow_frame, text="▼", command=move_root_down, bg="#2196F3", fg="white", font=("Arial", 12, "bold"), 
          relief="flat", bd=0, width=3, height=1).pack(pady=2)

# Action buttons frame on the right
roots_btn_frame = tk.Frame(roots_content_frame)
roots_btn_frame.pack(side="right", fill="y")
tk.Button(roots_btn_frame, text="➕ Add Root", command=add_root_gui, bg="#4CAF50", fg="white", font=("Arial", 9, "bold"), relief="flat", bd=0).pack(fill="x", pady=2, padx=2)
tk.Button(roots_btn_frame, text="➖ Remove Root", command=remove_root_gui, bg="#f44336", fg="white", font=("Arial", 9, "bold"), relief="flat", bd=0).pack(fill="x", pady=2, padx=2)

update_roots_listbox()

# Search Controls Frame (left, below roots)
search_frame = tk.LabelFrame(root, text="🔍 Search", font=("Arial", 11, "bold"), fg="navy", relief="groove", bd=2)
search_frame.grid(row=1, column=0, padx=10, pady=5, sticky="ew")
search_inner = tk.Frame(search_frame)
search_inner.pack(fill="x", expand=True, padx=5, pady=5)
keyword_label = tk.Label(search_inner, text="Keyword:", font=("Arial", 10, "bold"))
keyword_label.pack(side="left", padx=5)
keyword_entry = tk.Entry(search_inner, width=50, font=("Arial", 10), relief="flat", bd=1)
keyword_entry.pack(side="left", padx=5)

# Search buttons frame
search_buttons_frame = tk.Frame(search_inner)
search_buttons_frame.pack(side="left", padx=5)

# Add AI model selection combobox
ai_model_var = tk.StringVar(value="local")
ai_model_options = ["local", "openai", "cohere"]
tk.Label(search_inner, text="AI Model:", font=("Arial", 9)).pack(side="left", padx=5)
ai_model_menu = ttk.Combobox(search_inner, textvariable=ai_model_var, values=ai_model_options, state="readonly", width=10)
ai_model_menu.pack(side="left", padx=2)

tk.Button(search_buttons_frame, text="🔍 Live Search", command=start_live_search_thread, bg="#2196F3", fg="white", font=("Arial", 9, "bold"), relief="flat", bd=0).pack(side="left", padx=2)
tk.Button(search_buttons_frame, text="⚡ Index Search", command=start_index_search, bg="#FF9800", fg="white", font=("Arial", 9, "bold"), relief="flat", bd=0).pack(side="left", padx=2)
tk.Button(search_buttons_frame, text="🤖 AI Search", command=start_ai_search, bg="#9C27B0", fg="white", font=("Arial", 9, "bold"), relief="flat", bd=0).pack(side="left", padx=2)
tk.Button(search_buttons_frame, text="🗑️ Clear", command=clear_results, bg="#9E9E9E", fg="white", font=("Arial", 9, "bold"), relief="flat", bd=0).pack(side="left", padx=2)

# Index Management Controls Frame (right, full height below roots)
index_frame = tk.LabelFrame(root, text="⚙️ Index Management", font=("Arial", 11, "bold"), fg="navy", relief="groove", bd=2)
index_frame.grid(row=1, column=1, rowspan=2, padx=10, pady=5, sticky="nsew")
index_inner = tk.Frame(index_frame)
index_inner.pack(fill="both", expand=True, padx=5, pady=5)

# Results Frame (left, bottom)
results_frame = tk.LabelFrame(root, text="Search Results", font=("Arial", 12, "bold"), fg="navy", relief="groove", bd=2)
results_frame.grid(row=2, column=0, padx=10, pady=10, sticky="nsew")

# Create a frame for the scrollable area
results_scroll_frame = tk.Frame(results_frame)
results_scroll_frame.pack(fill="both", expand=True, padx=5, pady=5)

canvas = tk.Canvas(results_scroll_frame, width=900, height=400, bg="white", relief="flat", bd=0)
scrollbar = tk.Scrollbar(results_scroll_frame, orient="vertical", command=canvas.yview)
results_inner_frame = tk.Frame(canvas, bg="white")
results_inner_frame.bind(
    "<Configure>",
    lambda e: canvas.configure(
        scrollregion=canvas.bbox("all")
    )
)
canvas.create_window((0, 0), window=results_inner_frame, anchor="nw")
canvas.configure(yscrollcommand=scrollbar.set)
canvas.pack(side="left", fill="both", expand=True)
scrollbar.pack(side="right", fill="y")
root.grid_rowconfigure(1, weight=1)
root.grid_rowconfigure(2, weight=2)
root.grid_columnconfigure(0, weight=3)
root.grid_columnconfigure(1, weight=1)

# Regular Index Section
regular_index_section = tk.LabelFrame(index_inner, text="Regular Index", font=("Arial", 9, "bold"), fg="#333", relief="ridge", bd=1)
regular_index_section.pack(fill="x", pady=2, anchor="w")
tk.Button(regular_index_section, text="🔄 Rebuild", command=build_index_all_thread, bg="#FF5722", fg="white", font=("Arial", 8, "bold"), relief="flat", bd=0, width=12).pack(side="left", padx=2, pady=2)
tk.Button(regular_index_section, text="🔄 Update", command=update_index_all_thread, bg="#607D8B", fg="white", font=("Arial", 8, "bold"), relief="flat", bd=0, width=12).pack(side="left", padx=2, pady=2)

# AI Index Section
ai_index_section = tk.LabelFrame(index_inner, text="AI Index", font=("Arial", 9, "bold"), fg="#333", relief="ridge", bd=1)
ai_index_section.pack(fill="x", pady=2, anchor="w")
tk.Button(ai_index_section, text="🤖 Build AI", command=build_ai_index, bg="#9C27B0", fg="white", font=("Arial", 8, "bold"), relief="flat", bd=0, width=16).pack(side="left", padx=2, pady=2)
tk.Button(ai_index_section, text="🗑️ Clear AI", command=clear_ai_index_gui, bg="#E91E63", fg="white", font=("Arial", 8, "bold"), relief="flat", bd=0, width=16).pack(side="left", padx=2, pady=2)

# External Embeddings Section
embedding_section = tk.LabelFrame(index_inner, text="External Embeddings", font=("Arial", 9, "bold"), fg="#333", relief="ridge", bd=1)
embedding_section.pack(fill="x", pady=2, anchor="w")

# === Threaded Embedding Build Functions ===
def build_openai_embeddings_thread():
    wait_win = show_wait_message("Building OpenAI embeddings, please wait...")
    def run():
        success = build_openai_embeddings()
        root.after(0, lambda: wait_win.destroy())
        if success:
            root.after(0, lambda: messagebox.showinfo("OpenAI Embeddings", "OpenAI embeddings built and saved!"))
        else:
            root.after(0, lambda: messagebox.showerror("OpenAI Embeddings", "Failed to build OpenAI embeddings."))
    threading.Thread(target=run).start()

def build_cohere_embeddings_thread():
    wait_win = show_wait_message("Building Cohere embeddings, please wait...")
    def run():
        success = build_cohere_embeddings()
        root.after(0, lambda: wait_win.destroy())
        if success:
            root.after(0, lambda: messagebox.showinfo("Cohere Embeddings", "Cohere embeddings built and saved!"))
        else:
            root.after(0, lambda: messagebox.showerror("Cohere Embeddings", "Failed to build Cohere embeddings."))
    threading.Thread(target=run).start()

# --- Add embedding build buttons (vertical stack for visibility) ---
tk.Button(embedding_section, text="🔗 Build OpenAI Embeddings", command=build_openai_embeddings_thread, bg="#1976D2", fg="white", font=("Arial", 8, "bold"), relief="flat", bd=0, width=24).pack(fill="x", padx=2, pady=2)
tk.Button(embedding_section, text="🔗 Build Cohere Embeddings", command=build_cohere_embeddings_thread, bg="#00BFAE", fg="white", font=("Arial", 8, "bold"), relief="flat", bd=0, width=24).pack(fill="x", padx=2, pady=2)

# === Embedding Build Functions ===
def build_openai_embeddings():
    from openai import OpenAI
    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        print("[OpenAI API key not set]")
        return False
    client = OpenAI(api_key=api_key)
    # Get all indexed files
    conn = sqlite3.connect(INDEX_DB, timeout=30)
    c = conn.cursor()
    c.execute('SELECT file_path, file_type, content FROM file_index')
    files = c.fetchall()
    conn.close()
    if not files:
        print("[No files found in index]")
        return False
    embeddings = {}
    for file_path, file_type, content in files:
        try:
            resp = client.embeddings.create(
                input=content[:2000],
                model="text-embedding-ada-002"
            )
            emb = resp.data[0].embedding
            embeddings[file_path] = emb
        except Exception as e:
            print(f"[OpenAI embedding failed for {file_path}: {e}]")
    with open("embeddings_openai.json", "w", encoding="utf-8") as f:
        json.dump(embeddings, f)
    print("[OpenAI embeddings built and saved]")
    return True

def build_cohere_embeddings():
    import cohere
    api_key = os.getenv("COHERE_API_KEY")
    if not api_key:
        print("[Cohere API key not set]")
        return False
    co = cohere.Client(api_key)
    # Get all indexed files
    conn = sqlite3.connect(INDEX_DB, timeout=30)
    c = conn.cursor()
    c.execute('SELECT file_path, file_type, content FROM file_index')
    files = c.fetchall()
    conn.close()
    if not files:
        print("[No files found in index]")
        return False
    embeddings = {}
    for file_path, file_type, content in files:
        try:
            resp = co.embed(texts=[content[:2000]], model="embed-english-v3.0", input_type="search_document")
            emb = resp.embeddings[0]
            embeddings[file_path] = emb
        except Exception as e:
            print(f"[Cohere embedding failed for {file_path}: {e}]")
    with open("embeddings_cohere.json", "w", encoding="utf-8") as f:
        json.dump(embeddings, f)
    print("[Cohere embeddings built and saved]")
    return True

# === AI Dispatch Functions ===
def ai_search_dispatch(keyword, n_results, model_choice):
    """Dispatch AI search to the selected backend."""
    if model_choice == "local":
        return ai_search(keyword, n_results=n_results)
    elif model_choice == "openai":
        return openai_ai_search(keyword, n_results)
    elif model_choice == "cohere":
        return cohere_ai_search(keyword, n_results)
    else:
        return []

def ai_summarize_dispatch(text, model_choice):
    """Dispatch summarization to the selected backend."""
    if model_choice == "local":
        try:
            from ai_search import ai_engine
            if hasattr(ai_engine, 'summarizer') and ai_engine.summarizer:
                return ai_engine.summarizer(text[:2048], max_length=200, min_length=50, do_sample=False)[0]['summary_text']
        except Exception as e:
            print(f"[DEBUG] Local summarizer failed: {e}")
        return text[:300]  # fallback
    elif model_choice == "openai":
        return openai_summarize(text)
    elif model_choice == "cohere":
        return cohere_summarize(text)
    else:
        return text[:300]

# === External AI API stubs ===
def openai_summarize(text):
    try:
        from openai import OpenAI
        api_key = os.getenv("OPENAI_API_KEY")
        if not api_key:
            return "[OpenAI API key not set]"
        client = OpenAI(api_key=api_key)
        response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "Summarize the following text."},
                {"role": "user", "content": text[:3000]}
            ],
            max_tokens=200,
            temperature=0.3,
        )
        result = response.choices[0].message.content
        if result is None:
            print(f"[OpenAI summarization failed: No content in response: {response}]")
            return "[OpenAI summarization failed: No summary returned]"
        return result.strip()
    except Exception as e:
        return f"[OpenAI summarization failed: {e}]"

def cohere_summarize(text):
    try:
        import cohere
        api_key = os.getenv("COHERE_API_KEY")
        if not api_key:
            return "[Cohere API key not set]"
        co = cohere.Client(api_key)
        response = co.summarize(text=text[:3000], model='summarize-xlarge', length='medium', format='paragraph')
        return response.summary
    except Exception as e:
        return f"[Cohere summarization failed: {e}]"

def openai_ai_search(keyword, n_results):
    from openai import OpenAI
    import numpy as np
    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        print("[OpenAI API key not set]")
        return []
    client = OpenAI(api_key=api_key)
    embeddings = load_embeddings("embeddings_openai.json")
    if not embeddings:
        print("[No OpenAI embeddings found. Run embedding build first]")
        return []
    # Embed the query
    try:
        resp = client.embeddings.create(input=keyword, model="text-embedding-ada-002")
        query_emb = np.array(resp.data[0].embedding)
    except Exception as e:
        print(f"[OpenAI query embedding failed: {e}]")
        return []
    # Compute similarities
    scored = []
    for file_path, emb in embeddings.items():
        emb_vec = np.array(emb)
        sim = np.dot(query_emb, emb_vec) / (np.linalg.norm(query_emb) * np.linalg.norm(emb_vec) + 1e-8)
        scored.append((file_path, sim))
    scored.sort(key=lambda x: x[1], reverse=True)
    # Get file info for top-N
    conn = sqlite3.connect(INDEX_DB, timeout=30)
    c = conn.cursor()
    results = []
    for file_path, sim in scored[:n_results]:
        c.execute('SELECT file_type, content FROM file_index WHERE file_path=?', (file_path,))
        row = c.fetchone()
        if row:
            file_type, content = row
            results.append({
                'file_path': file_path,
                'file_type': file_type,
                'similarity_score': sim,
                'content': content
            })
    conn.close()
    return results

def cohere_ai_search(keyword, n_results):
    import cohere
    import numpy as np
    api_key = os.getenv("COHERE_API_KEY")
    if not api_key:
        print("[Cohere API key not set]")
        return []
    co = cohere.Client(api_key)
    embeddings = load_embeddings("embeddings_cohere.json")
    if not embeddings:
        print("[No Cohere embeddings found. Run embedding build first]")
        return []
    # Embed the query
    try:
        resp = co.embed(texts=[keyword], model="embed-english-v3.0", input_type="search_query")
        query_emb = np.array(resp.embeddings[0])
    except Exception as e:
        print(f"[Cohere query embedding failed: {e}]")
        return []
    # Compute similarities
    scored = []
    for file_path, emb in embeddings.items():
        emb_vec = np.array(emb)
        sim = np.dot(query_emb, emb_vec) / (np.linalg.norm(query_emb) * np.linalg.norm(emb_vec) + 1e-8)
        scored.append((file_path, sim))
    scored.sort(key=lambda x: x[1], reverse=True)
    # Get file info for top-N
    conn = sqlite3.connect(INDEX_DB, timeout=30)
    c = conn.cursor()
    results = []
    for file_path, sim in scored[:n_results]:
        c.execute('SELECT file_type, content FROM file_index WHERE file_path=?', (file_path,))
        row = c.fetchone()
        if row:
            file_type, content = row
            results.append({
                'file_path': file_path,
                'file_type': file_type,
                'similarity_score': sim,
                'content': content
            })
    conn.close()
    return results

root.mainloop()