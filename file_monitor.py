#!/usr/bin/env python3
"""
File Monitor and Auto-Converter
Monitors folders for new or changed files and automatically converts them using existing converters.
"""

import os
import time
import threading
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
from pathlib import Path
import shutil
import hashlib
import json
from datetime import datetime
import subprocess
import sys
import re
from docx import Document
import PyPDF2
from pdfminer.high_level import extract_text as pdfminer_extract_text

# Standalone conversion functions (no GUI windows)
def convert_doc_to_docx_standalone(doc_path, docx_path):
    """Convert DOC to DOCX without GUI"""
    try:
        # Normalize paths and handle encoding issues
        doc_path = os.path.abspath(doc_path)
        docx_path = os.path.abspath(docx_path)
        
        # Check if source file exists
        if not os.path.exists(doc_path):
            print(f"Source file not found: {doc_path}")
            return False
            
        # Skip temporary files (files starting with ~$)
        if os.path.basename(doc_path).startswith('~$'):
            print(f"Skipping temporary file: {os.path.basename(doc_path)}")
            return False
            
        # Create output directory if needed
        os.makedirs(os.path.dirname(docx_path), exist_ok=True)
        
        print(f"Converting: {os.path.basename(doc_path)}")
        
        # Method 1: Try using LibreOffice (if available)
        print("  Trying LibreOffice conversion...")
        if try_libreoffice_conversion(doc_path, docx_path):
            print("  ✓ LibreOffice conversion successful")
            return True
            
        # Method 2: Try using Microsoft Word (if available)
        print("  Trying Microsoft Word conversion...")
        if try_word_conversion(doc_path, docx_path):
            print("  ✓ Microsoft Word conversion successful")
            return True
            
        # Method 3: Try to copy if it's already a DOCX in disguise
        print("  Checking if file is already DOCX format...")
        if try_copy_if_docx(doc_path, docx_path):
            print("  ✓ File was already DOCX format")
            return True
            
        # Method 4: Try using python-docx (limited success)
        print("  Trying python-docx conversion...")
        if try_python_docx_conversion(doc_path, docx_path):
            print("  ✓ python-docx conversion successful")
            return True
            
        print("  ✗ All conversion methods failed")
        return False
        
    except Exception as e:
        print(f"Error converting {doc_path}: {e}")
        return False

def try_libreoffice_conversion(doc_path, docx_path):
    """Try conversion using LibreOffice"""
    try:
        # Check if LibreOffice is available
        libreoffice_paths = [
            "soffice",  # Linux/Mac
            "libreoffice",  # Linux
            r"C:\Program Files\LibreOffice\program\soffice.exe",  # Windows
            r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",  # Windows 32-bit
            r"C:\Program Files\LibreOffice*\program\soffice.exe",  # Windows with version
            r"C:\Program Files (x86)\LibreOffice*\program\soffice.exe",  # Windows 32-bit with version
        ]
        
        for path in libreoffice_paths:
            try:
                # Handle wildcards in paths
                if '*' in path:
                    import glob
                    matching_paths = glob.glob(path)
                    if not matching_paths:
                        continue
                    path = matching_paths[0]  # Use the first match
                
                # Test if the executable exists
                if not os.path.exists(path) and not path in ["soffice", "libreoffice"]:
                    continue
                
                cmd = [path, "--headless", "--convert-to", "docx", 
                       "--outdir", os.path.dirname(docx_path), doc_path]
                result = subprocess.run(cmd, capture_output=True, text=True, timeout=60)
                
                if result.returncode == 0:
                    # LibreOffice creates the file with a different name
                    expected_name = os.path.splitext(os.path.basename(doc_path))[0] + ".docx"
                    expected_path = os.path.join(os.path.dirname(docx_path), expected_name)
                    
                    if os.path.exists(expected_path):
                        if expected_path != docx_path:
                            shutil.move(expected_path, docx_path)
                        return True
                else:
                    print(f"    LibreOffice command failed with return code: {result.returncode}")
                    if result.stderr:
                        print(f"    Error: {result.stderr.strip()}")
            except Exception as e:
                print(f"    LibreOffice path {path} failed: {e}")
                continue
                
        print("    LibreOffice not found or conversion failed")
        return False
        
    except Exception as e:
        print(f"LibreOffice conversion failed: {e}")
        return False

def explain_doc_conversion_issues():
    """Explain common reasons why DOC to DOCX conversion fails"""
    print("\n🔍 Common reasons why DOC to DOCX conversion fails:")
    print("1. LibreOffice not installed - Install LibreOffice for best results")
    print("2. Microsoft Word not available - Install Microsoft Word for Windows")
    print("3. Corrupted DOC files - Old or damaged DOC files may not convert")
    print("4. Complex formatting - Some advanced formatting may not convert properly")
    print("5. File permissions - Ensure read/write access to input/output folders")
    print("6. File size too large - Very large files may timeout")
    print("\n💡 Solutions:")
    print("- Install LibreOffice: https://www.libreoffice.org/")
    print("- Install Microsoft Word (Windows only)")
    print("- Try converting files manually first to test")
    print("- Check file permissions and disk space")

def try_word_conversion(doc_path, docx_path):
    """Try conversion using Microsoft Word (Windows only)"""
    try:
        import win32com.client
        
        # Convert path to absolute path and fix slashes
        abs_doc_path = os.path.abspath(doc_path).replace('/', '\\')
        abs_docx_path = os.path.abspath(docx_path).replace('/', '\\')
        
        # Check if file exists
        if not os.path.exists(abs_doc_path):
            print(f"File not found: {abs_doc_path}")
            return False
        
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        
        try:
            doc = word.Documents.Open(abs_doc_path)
            doc.SaveAs2(abs_docx_path, FileFormat=16)  # 16 = docx format
            doc.Close()
            return True
        except Exception as e:
            print(f"Word document operation failed: {e}")
            return False
        finally:
            word.Quit()
        
    except Exception as e:
        print(f"Word conversion failed: {e}")
        return False

def try_python_docx_conversion(doc_path, docx_path):
    """Try conversion using python-docx (limited success)"""
    try:
        # Check if file exists and is readable
        if not os.path.exists(doc_path):
            print(f"File not found: {doc_path}")
            return False
            
        # Try to read the file first
        try:
            with open(doc_path, 'rb') as f:
                # Check if it's actually a DOCX file in disguise
                content = f.read(4)
                if content == b'PK\x03\x04':
                    # It's actually a DOCX file, just copy it
                    shutil.copy2(doc_path, docx_path)
                    return True
        except Exception as e:
            print(f"Cannot read file {doc_path}: {e}")
            return False
        
        # This method has limited success with old DOC files
        # It works better with newer DOC files that are actually DOCX in disguise
        doc = Document(doc_path)
        doc.save(docx_path)
        return True
        
    except Exception as e:
        print(f"python-docx conversion failed: {e}")
        return False

def try_copy_if_docx(doc_path, docx_path):
    """Check if file is actually a DOCX file in disguise and copy it"""
    try:
        with open(doc_path, 'rb') as f:
            # Check for ZIP file signature (DOCX files are ZIP archives)
            content = f.read(4)
            if content == b'PK\x03\x04':
                print(f"File {os.path.basename(doc_path)} is actually a DOCX file, copying...")
                shutil.copy2(doc_path, docx_path)
                return True
        return False
    except Exception as e:
        print(f"Error checking file format: {e}")
        return False

def convert_docx_to_markdown_standalone(docx_path, md_path):
    """Convert DOCX to Markdown without GUI"""
    try:
        doc = Document(docx_path)
        markdown_content = []
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                markdown_content.append("")
                continue
                
            # Check for headings
            style_name = paragraph.style.name if paragraph.style and paragraph.style.name else ""
            if style_name.startswith('Heading'):
                level_str = style_name[-1] if style_name[-1].isdigit() else "1"
                level = int(level_str) if level_str.isdigit() else 1
                markdown_content.append(f"{'#' * level} {text}")
            else:
                # Check for bold and italic
                formatted_text = text
                for run in paragraph.runs:
                    if run.bold and run.italic:
                        formatted_text = formatted_text.replace(run.text, f"***{run.text}***")
                    elif run.bold:
                        formatted_text = formatted_text.replace(run.text, f"**{run.text}**")
                    elif run.italic:
                        formatted_text = formatted_text.replace(run.text, f"*{run.text}*")
                
                markdown_content.append(formatted_text)
                
        # Write markdown file
        with open(md_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(markdown_content))
            
        return True
        
    except Exception as e:
        print(f"Error converting DOCX {docx_path}: {e}")
        return False

class FileMonitor:
    def __init__(self):
        self.root = tk.Tk()
        self.root.title("📁 File Monitor & Auto-Converter")
        self.root.geometry("900x700")
        self.root.configure(bg="#f5f5f5")
        
        # Variables
        self.monitor_folders = []
        self.docx_output_folder = ""
        self.markdown_output_folder = ""
        self.monitor_interval = tk.IntVar(value=2)  # Default 2 seconds
        self.is_monitoring = False
        self.monitor_thread = None
        self.file_hashes = {}  # Store file hashes to detect changes
        self.config_file = "file_monitor_config.json"
        self.progress_file = "file_monitor_progress.json"  # Track processed files
        self.processed_files = {}  # Dictionary to track processed files and their status
        
        # Load configuration and progress
        self.load_config()
        self.load_progress()
        
        self.setup_ui()
        
    def setup_ui(self):
        """Setup the user interface"""
        # Main frame
        main_frame = tk.Frame(self.root, bg="#f5f5f5")
        main_frame.pack(fill="both", expand=True, padx=20, pady=20)
        
        # Title
        title_label = tk.Label(main_frame, text="📁 File Monitor & Auto-Converter", 
                             font=("Arial", 16, "bold"), bg="#f5f5f5", fg="#2196F3")
        title_label.pack(pady=(0, 20))
        
        # Monitor folders frame
        folders_frame = tk.LabelFrame(main_frame, text="📁 Monitor Folders", 
                                    font=("Arial", 11, "bold"), bg="#f5f5f5", fg="navy")
        folders_frame.pack(fill="x", pady=10)
        
        folders_inner = tk.Frame(folders_frame, bg="#f5f5f5")
        folders_inner.pack(fill="x", padx=10, pady=10)
        
        # Folders listbox
        self.folders_listbox = tk.Listbox(folders_inner, height=4, font=("Arial", 9))
        self.folders_listbox.pack(side="left", fill="x", expand=True, padx=(0, 10))
        
        # Folders buttons
        folders_btn_frame = tk.Frame(folders_inner, bg="#f5f5f5")
        folders_btn_frame.pack(side="right", fill="y")
        
        tk.Button(folders_btn_frame, text="➕ Add Folder", command=self.add_monitor_folder, 
                 bg="#4CAF50", fg="white", font=("Arial", 9, "bold"), relief="flat", bd=0).pack(fill="x", pady=2)
        tk.Button(folders_btn_frame, text="➖ Remove Folder", command=self.remove_monitor_folder, 
                 bg="#f44336", fg="white", font=("Arial", 9, "bold"), relief="flat", bd=0).pack(fill="x", pady=2)
        
        # Output folders frame
        output_frame = tk.LabelFrame(main_frame, text="📁 Output Folders", 
                                   font=("Arial", 11, "bold"), bg="#f5f5f5", fg="navy")
        output_frame.pack(fill="x", pady=10)
        
        # DOCX output
        docx_frame = tk.Frame(output_frame, bg="#f5f5f5")
        docx_frame.pack(fill="x", padx=10, pady=5)
        tk.Label(docx_frame, text="DOCX Output:", font=("Arial", 9, "bold"), bg="#f5f5f5").pack(side="left")
        self.docx_entry = tk.Entry(docx_frame, textvariable=tk.StringVar(value=self.docx_output_folder), 
                width=50, font=("Arial", 9))
        self.docx_entry.pack(side="left", padx=5, fill="x", expand=True)
        tk.Button(docx_frame, text="Browse", command=self.browse_docx_output, 
                 bg="#2196F3", fg="white", font=("Arial", 8, "bold"), relief="flat", bd=0).pack(side="right")
        
        # Markdown output
        md_frame = tk.Frame(output_frame, bg="#f5f5f5")
        md_frame.pack(fill="x", padx=10, pady=5)
        tk.Label(md_frame, text="Markdown Output:", font=("Arial", 9, "bold"), bg="#f5f5f5").pack(side="left")
        self.md_entry = tk.Entry(md_frame, textvariable=tk.StringVar(value=self.markdown_output_folder), 
                width=50, font=("Arial", 9))
        self.md_entry.pack(side="left", padx=5, fill="x", expand=True)
        tk.Button(md_frame, text="Browse", command=self.browse_markdown_output, 
                 bg="#2196F3", fg="white", font=("Arial", 8, "bold"), relief="flat", bd=0).pack(side="right")
        
        # Status frame
        status_frame = tk.LabelFrame(main_frame, text="📊 Monitor Status", 
                                   font=("Arial", 11, "bold"), bg="#f5f5f5", fg="navy")
        status_frame.pack(fill="x", pady=10)
        
        status_inner = tk.Frame(status_frame, bg="#f5f5f5")
        status_inner.pack(fill="x", padx=10, pady=10)
        
        # Status label
        self.status_label = tk.Label(status_inner, text="⏸️ Monitoring stopped", 
                                   font=("Arial", 10, "bold"), bg="#f5f5f5", fg="#666666")
        self.status_label.pack()
        
        # Monitoring interval control
        interval_frame = tk.Frame(status_inner, bg="#f5f5f5")
        interval_frame.pack(fill="x", pady=(10, 0))
        
        tk.Label(interval_frame, text="Check interval (seconds):", 
                font=("Arial", 9), bg="#f5f5f5").pack(side="left")
        
        interval_spinbox = tk.Spinbox(interval_frame, from_=1, to=60, width=5, 
                                     textvariable=self.monitor_interval, 
                                     font=("Arial", 9))
        interval_spinbox.pack(side="left", padx=(5, 0))
        
        tk.Label(interval_frame, text="(1-60 seconds)", 
                font=("Arial", 8), fg="#666666", bg="#f5f5f5").pack(side="left", padx=(5, 0))
        
        # Log frame
        log_frame = tk.LabelFrame(main_frame, text="📝 Activity Log", 
                                font=("Arial", 11, "bold"), bg="#f5f5f5", fg="navy")
        log_frame.pack(fill="both", expand=True, pady=10)
        
        log_inner = tk.Frame(log_frame, bg="#f5f5f5")
        log_inner.pack(fill="both", expand=True, padx=10, pady=10)
        
        # Create text widget with scrollbar
        self.log_text = tk.Text(log_inner, height=10, font=("Consolas", 9), bg="white", fg="black")
        scrollbar = tk.Scrollbar(log_inner, orient="vertical", command=self.log_text.yview)
        self.log_text.configure(yscrollcommand=scrollbar.set)
        
        self.log_text.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")
        
        # Control buttons frame
        buttons_frame = tk.Frame(main_frame, bg="#f5f5f5")
        buttons_frame.pack(fill="x", pady=(20, 0))
        
        button_sub_frame = tk.Frame(buttons_frame, bg="#f5f5f5")
        button_sub_frame.pack(expand=True)
        
        self.start_button = tk.Button(button_sub_frame, text="▶️ Start Monitoring", command=self.start_monitoring, 
                                    bg="#4CAF50", fg="white", font=("Arial", 12, "bold"), 
                                    relief="flat", bd=0, height=2, width=15)
        self.start_button.pack(side="left", padx=5)
        
        self.stop_button = tk.Button(button_sub_frame, text="⏸️ Stop Monitoring", command=self.stop_monitoring, 
                                   bg="#f44336", fg="white", font=("Arial", 12, "bold"), 
                                   relief="flat", bd=0, height=2, width=15, state="disabled")
        self.stop_button.pack(side="left", padx=5)
        
        tk.Button(button_sub_frame, text="🗑️ Clear Log", command=self.clear_log, 
                 bg="#FF9800", fg="white", font=("Arial", 12, "bold"), 
                 relief="flat", bd=0, height=2, width=15).pack(side="left", padx=5)
        
        tk.Button(button_sub_frame, text="📊 Progress", command=self.show_progress_stats, 
                 bg="#2196F3", fg="white", font=("Arial", 12, "bold"), 
                 relief="flat", bd=0, height=2, width=15).pack(side="left", padx=5)
        
        tk.Button(button_sub_frame, text="🔄 Reset Progress", command=self.clear_progress, 
                 bg="#E91E63", fg="white", font=("Arial", 12, "bold"), 
                 relief="flat", bd=0, height=2, width=15).pack(side="left", padx=5)
        
        tk.Button(button_sub_frame, text="❌ Exit", command=self.exit_app, 
                 bg="#9E9E9E", fg="white", font=("Arial", 12, "bold"), 
                 relief="flat", bd=0, height=2, width=15).pack(side="right", padx=5)
        
        # Update folders listbox
        self.update_folders_listbox()
        
        # Show initial progress status
        total, completed, failed = self.get_progress_stats()
        if total > 0:
            self.log_message(f"📊 Loaded progress: {completed} completed, {failed} failed, {total} total files tracked")
        else:
            self.log_message("📊 No previous progress found - starting fresh")
        
    def add_monitor_folder(self):
        """Add a folder to monitor"""
        folder = filedialog.askdirectory(title="Select folder to monitor")
        if folder and folder not in self.monitor_folders:
            self.monitor_folders.append(folder)
            self.update_folders_listbox()
            self.save_config()
            self.log_message(f"Added monitor folder: {folder}")
            
            # Clear progress when adding new monitor folders
            if self.processed_files:
                self.log_message("⚠️ Monitor folders changed - clearing progress tracking")
                self.clear_progress()
            
    def remove_monitor_folder(self):
        """Remove selected folder from monitoring"""
        selection = self.folders_listbox.curselection()
        if selection:
            folder = self.folders_listbox.get(selection[0])
            self.monitor_folders.remove(folder)
            self.update_folders_listbox()
            self.save_config()
            self.log_message(f"Removed monitor folder: {folder}")
            
            # Clear progress when removing monitor folders
            if self.processed_files:
                self.log_message("⚠️ Monitor folders changed - clearing progress tracking")
                self.clear_progress()
            
    def update_folders_listbox(self):
        """Update the folders listbox"""
        self.folders_listbox.delete(0, tk.END)
        for folder in self.monitor_folders:
            self.folders_listbox.insert(tk.END, folder)
            
    def browse_docx_output(self):
        """Browse for DOCX output folder"""
        folder = filedialog.askdirectory(title="Select folder for DOCX files")
        if folder:
            old_folder = self.docx_output_folder
            self.docx_output_folder = folder
            self.docx_entry.delete(0, tk.END)
            self.docx_entry.insert(0, folder)
            self.save_config()
            self.log_message(f"Set DOCX output folder: {folder}")
            
            # Clear progress if output folder changed
            if old_folder and old_folder != folder:
                self.log_message("⚠️ Output folder changed - clearing progress tracking")
                self.clear_progress()
            
    def browse_markdown_output(self):
        """Browse for Markdown output folder"""
        folder = filedialog.askdirectory(title="Select folder for Markdown files")
        if folder:
            old_folder = self.markdown_output_folder
            self.markdown_output_folder = folder
            self.md_entry.delete(0, tk.END)
            self.md_entry.insert(0, folder)
            self.save_config()
            self.log_message(f"Set Markdown output folder: {folder}")
            
            # Clear progress if output folder changed
            if old_folder and old_folder != folder:
                self.log_message("⚠️ Output folder changed - clearing progress tracking")
                self.clear_progress()
            
    def start_monitoring(self):
        """Start monitoring folders"""
        if not self.monitor_folders:
            messagebox.showwarning("Warning", "Please add at least one folder to monitor.")
            return
            
        if not self.docx_output_folder or not self.markdown_output_folder:
            messagebox.showwarning("Warning", "Please set both DOCX and Markdown output folders.")
            return
            
        # Start real-time monitoring immediately
        self.is_monitoring = True
        self.monitor_thread = threading.Thread(target=self.monitor_loop, daemon=True)
        self.monitor_thread.start()
        
        self.start_button.config(state="disabled")
        self.stop_button.config(state="normal")
        self.status_label.config(text="🟢 Monitoring active", fg="#4CAF50")
        self.log_message("Started real-time monitoring")
        
        # Process existing files in background thread
        self.log_message("Processing existing files in background...")
        initial_process_thread = threading.Thread(target=self.process_existing_files, daemon=True)
        initial_process_thread.start()
        
    def process_existing_files(self):
        """Process all existing .doc and .docx files in background thread"""
        try:
            total_files = 0
            processed_files = 0
            skipped_files = 0
            
            for folder in self.monitor_folders:
                for root, dirs, files in os.walk(folder):
                    for file in files:
                        file_path = os.path.join(root, file)
                        file_ext = os.path.splitext(file)[1].lower()
                        
                        if file_ext in ['.doc', '.docx', '.pdf', '.txt']:
                            total_files += 1
                            
                            # Check if file is already processed
                            if self.is_file_processed(file_path):
                                self.log_message(f"⏭️ Skipping already processed: {os.path.basename(file_path)}")
                                skipped_files += 1
                                self.file_hashes[file_path] = self.get_file_hash(file_path)
                                continue
                            
                            # Process the file
                            if file_ext == '.doc':
                                self.log_message(f"🔄 Processing DOC file: {os.path.basename(file_path)}")
                                if self.convert_doc_to_docx_and_markdown(file_path):
                                    self.mark_file_processed(file_path)
                                    processed_files += 1
                                else:
                                    self.mark_file_failed(file_path, "DOC to DOCX/MD conversion failed")
                            elif file_ext == '.docx':
                                self.log_message(f"🔄 Processing DOCX file: {os.path.basename(file_path)}")
                                if self.convert_docx_to_markdown(file_path):
                                    self.mark_file_processed(file_path)
                                    processed_files += 1
                                else:
                                    self.mark_file_failed(file_path, "DOCX to MD conversion failed")
                            elif file_ext in ['.pdf', '.txt']:
                                self.log_message(f"🔄 Processing {file_ext.upper()} file: {os.path.basename(file_path)}")
                                if self.convert_to_markdown(file_path):
                                    self.mark_file_processed(file_path)
                                    processed_files += 1
                                else:
                                    self.mark_file_failed(file_path, f"{file_ext.upper()} to MD conversion failed")
                            
                            self.file_hashes[file_path] = self.get_file_hash(file_path)
            
            self.log_message(f"✓ Initial file processing complete!")
            self.log_message(f"📊 Summary: {processed_files} processed, {skipped_files} skipped, {total_files} total")
            
        except Exception as e:
            self.log_message(f"✗ Error during initial file processing: {e}")
        
    def stop_monitoring(self):
        """Stop monitoring folders"""
        self.is_monitoring = False
        self.start_button.config(state="normal")
        self.stop_button.config(state="disabled")
        self.status_label.config(text="⏸️ Monitoring stopped", fg="#666666")
        self.log_message("Stopped monitoring folders")
        
    def monitor_loop(self):
        """Main monitoring loop"""
        while self.is_monitoring:
            try:
                self.check_for_changes()
                time.sleep(self.monitor_interval.get())  # Check every interval seconds
            except Exception as e:
                self.log_message(f"Error in monitor loop: {e}")
                time.sleep(5)  # Wait longer on error
                
    def check_for_changes(self):
        """Check for new or changed files in monitored folders"""
        for folder in self.monitor_folders:
            if not os.path.exists(folder):
                continue
                
            for root, dirs, files in os.walk(folder):
                for file in files:
                    file_path = os.path.join(root, file)
                    file_ext = os.path.splitext(file)[1].lower()
                    
                    # Only process supported file types
                    if file_ext not in ['.doc', '.docx', '.pdf', '.txt']:
                        continue
                        
                    file_hash = self.get_file_hash(file_path)
                    
                    # Check if file is new or changed
                    if file_path not in self.file_hashes or self.file_hashes[file_path] != file_hash:
                        self.file_hashes[file_path] = file_hash
                        
                        # Check if file is already processed (for new files)
                        if self.is_file_processed(file_path):
                            self.log_message(f"⏭️ Skipping already processed: {os.path.basename(file_path)}")
                            continue
                            
                        self.process_file(file_path)
                        
    def get_file_hash(self, file_path):
        """Get hash of file for change detection"""
        try:
            stat = os.stat(file_path)
            # Use modification time and file size for quick hash
            return f"{stat.st_mtime}_{stat.st_size}"
        except:
            return "0"
            
    def process_file(self, file_path):
        """Process a new or changed file"""
        file_ext = os.path.splitext(file_path)[1].lower()
        
        if file_ext == '.doc':
            self.log_message(f"🔄 Processing DOC file: {os.path.basename(file_path)}")
            if self.convert_doc_to_docx_and_markdown(file_path):
                self.mark_file_processed(file_path)
            else:
                self.mark_file_failed(file_path, "DOC to DOCX/MD conversion failed")
        elif file_ext in ['.docx', '.pdf', '.txt']:
            self.log_message(f"🔄 Processing {file_ext.upper()} file: {os.path.basename(file_path)}")
            if self.convert_to_markdown(file_path):
                self.mark_file_processed(file_path)
            else:
                self.mark_file_failed(file_path, f"{file_ext.upper()} to MD conversion failed")
            
    def convert_doc_to_docx_and_markdown(self, doc_path):
        """Convert DOC to DOCX and then to Markdown"""
        try:
            # First convert DOC to DOCX
            docx_path = self.get_docx_path(doc_path)
            
            # Calculate markdown path based on original DOC file (not intermediate DOCX)
            md_path = self.get_markdown_path_from_original(doc_path)
            
            # Log the folder structure preservation
            self.log_message(f"Converting: {doc_path}")
            self.log_message(f"DOCX output: {docx_path}")
            self.log_message(f"Markdown output: {md_path}")
            
            if self.convert_doc_to_docx(doc_path, docx_path):
                self.log_message(f"✓ Converted DOC to DOCX: {os.path.basename(docx_path)}")
                
                # Then convert DOCX to Markdown
                if self.convert_docx_to_markdown_file(docx_path, md_path):
                    self.log_message(f"✓ Converted DOCX to Markdown: {os.path.basename(md_path)}")
                    return True
                else:
                    self.log_message(f"✗ Failed to convert DOCX to Markdown: {os.path.basename(docx_path)}")
                    return False
            else:
                self.log_message(f"✗ Failed to convert DOC to DOCX: {os.path.basename(doc_path)}")
                self.log_message("💡 Tip: Install LibreOffice or Microsoft Word for better DOC conversion")
                return False
                
        except Exception as e:
            self.log_message(f"✗ Error processing DOC file {os.path.basename(doc_path)}: {e}")
            return False
            
    def convert_to_markdown(self, source_path):
        """Convert DOCX, PDF, or TXT to Markdown"""
        try:
            file_ext = os.path.splitext(source_path)[1].lower()
            md_path = self.get_markdown_path(source_path)
            
            if file_ext == '.docx':
                if self.convert_docx_to_markdown_file(source_path, md_path):
                    self.log_message(f"✓ Converted DOCX to Markdown: {os.path.basename(md_path)}")
                    return True
                else:
                    self.log_message(f"✗ Failed to convert DOCX to Markdown: {os.path.basename(source_path)}")
                    return False
            elif file_ext == '.pdf':
                if self.convert_pdf_to_markdown_file(source_path, md_path):
                    self.log_message(f"✓ Converted PDF to Markdown: {os.path.basename(md_path)}")
                    return True
                else:
                    self.log_message(f"✗ Failed to convert PDF to Markdown: {os.path.basename(source_path)}")
                    return False
            elif file_ext == '.txt':
                if self.convert_txt_to_markdown_file(source_path, md_path):
                    self.log_message(f"✓ Converted TXT to Markdown: {os.path.basename(md_path)}")
                    return True
                else:
                    self.log_message(f"✗ Failed to convert TXT to Markdown: {os.path.basename(source_path)}")
                    return False
        except Exception as e:
            self.log_message(f"✗ Error processing {os.path.basename(source_path)}: {e}")
            return False

    def convert_docx_to_markdown(self, docx_path):
        """Convert DOCX to Markdown"""
        try:
            md_path = self.get_markdown_path(docx_path)
            if self.convert_docx_to_markdown_file(docx_path, md_path):
                self.log_message(f"✓ Converted DOCX to Markdown: {os.path.basename(md_path)}")
                return True
            else:
                self.log_message(f"✗ Failed to convert DOCX to Markdown: {os.path.basename(docx_path)}")
                return False
        except Exception as e:
            self.log_message(f"✗ Error processing DOCX file {os.path.basename(docx_path)}: {e}")
            return False
            
    def get_docx_path(self, doc_path):
        """Get the DOCX output path for a DOC file, preserving folder structure including root folder"""
        # Find which monitored folder contains this file
        for monitor_folder in self.monitor_folders:
            if doc_path.startswith(monitor_folder):
                # Get the root folder name (the monitored folder name)
                root_folder_name = os.path.basename(monitor_folder)
                if not root_folder_name:
                    root_folder_name = os.path.basename(os.path.dirname(monitor_folder))
                
                # Calculate relative path from monitored folder
                rel_path = os.path.relpath(doc_path, monitor_folder)
                # Get the directory structure
                rel_dir = os.path.dirname(rel_path)
                filename = os.path.splitext(os.path.basename(doc_path))[0] + '.docx'
                
                # Create output path preserving the full folder structure including root
                if rel_dir == '.':
                    # File is directly in the monitored folder
                    output_dir = os.path.join(self.docx_output_folder, root_folder_name)
                else:
                    # File is in a subfolder, preserve the structure with root folder
                    output_dir = os.path.join(self.docx_output_folder, root_folder_name, rel_dir)
                
                return os.path.join(output_dir, filename)
        
        # Fallback if file is not in any monitored folder
        filename = os.path.splitext(os.path.basename(doc_path))[0] + '.docx'
        return os.path.join(self.docx_output_folder, filename)
        
    def get_markdown_path_from_original(self, original_path):
        """Get the Markdown output path for a file, preserving folder structure from original path including root folder"""
        # Find which monitored folder contains this file
        for monitor_folder in self.monitor_folders:
            if original_path.startswith(monitor_folder):
                # Get the root folder name (the monitored folder name)
                root_folder_name = os.path.basename(monitor_folder)
                if not root_folder_name:
                    root_folder_name = os.path.basename(os.path.dirname(monitor_folder))
                
                # Calculate relative path from monitored folder
                rel_path = os.path.relpath(original_path, monitor_folder)
                # Get the directory structure
                rel_dir = os.path.dirname(rel_path)
                filename = os.path.splitext(os.path.basename(original_path))[0] + '.md'
                
                # Create output path preserving the full folder structure including root
                if rel_dir == '.':
                    # File is directly in the monitored folder
                    output_dir = os.path.join(self.markdown_output_folder, root_folder_name)
                else:
                    # File is in a subfolder, preserve the structure with root folder
                    output_dir = os.path.join(self.markdown_output_folder, root_folder_name, rel_dir)
                
                return os.path.join(output_dir, filename)
        
        # Fallback if file is not in any monitored folder
        filename = os.path.splitext(os.path.basename(original_path))[0] + '.md'
        return os.path.join(self.markdown_output_folder, filename)
        
    def get_markdown_path(self, source_path):
        """Get the Markdown output path for a file, preserving folder structure including root folder"""
        # Find which monitored folder contains this file
        for monitor_folder in self.monitor_folders:
            if source_path.startswith(monitor_folder):
                # Get the root folder name (the monitored folder name)
                root_folder_name = os.path.basename(monitor_folder)
                if not root_folder_name:
                    root_folder_name = os.path.basename(os.path.dirname(monitor_folder))
                
                # Calculate relative path from monitored folder
                rel_path = os.path.relpath(source_path, monitor_folder)
                # Get the directory structure
                rel_dir = os.path.dirname(rel_path)
                filename = os.path.splitext(os.path.basename(source_path))[0] + '.md'
                
                # Create output path preserving the full folder structure including root
                if rel_dir == '.':
                    # File is directly in the monitored folder
                    output_dir = os.path.join(self.markdown_output_folder, root_folder_name)
                else:
                    # File is in a subfolder, preserve the structure with root folder
                    output_dir = os.path.join(self.markdown_output_folder, root_folder_name, rel_dir)
                
                return os.path.join(output_dir, filename)
        
        # Fallback if file is not in any monitored folder
        filename = os.path.splitext(os.path.basename(source_path))[0] + '.md'
        return os.path.join(self.markdown_output_folder, filename)
        
    def convert_doc_to_docx(self, doc_path, docx_path):
        """Convert DOC to DOCX using standalone function"""
        try:
            # Create output directory if needed
            output_dir = os.path.dirname(docx_path)
            if output_dir:
                os.makedirs(output_dir, exist_ok=True)
                self.log_message(f"Created directory: {output_dir}")
            
            # Use the standalone conversion function
            return convert_doc_to_docx_standalone(doc_path, docx_path)
            
        except Exception as e:
            self.log_message(f"Error in DOC to DOCX conversion: {e}")
            return False
            
    def convert_docx_to_markdown_file(self, docx_path, md_path):
        """Convert DOCX to Markdown using standalone function"""
        try:
            # Create output directory if needed
            output_dir = os.path.dirname(md_path)
            if output_dir:
                os.makedirs(output_dir, exist_ok=True)
                self.log_message(f"Created directory: {output_dir}")
            
            # Use the standalone conversion function
            return convert_docx_to_markdown_standalone(docx_path, md_path)
            
        except Exception as e:
            self.log_message(f"Error in DOCX to Markdown conversion: {e}")
            return False

    def convert_pdf_to_markdown_file(self, pdf_path, md_path):
        """Convert PDF to Markdown"""
        try:
            # Create output directory if needed
            output_dir = os.path.dirname(md_path)
            if output_dir:
                os.makedirs(output_dir, exist_ok=True)
                self.log_message(f"Created directory: {output_dir}")
            
            # Extract text from PDF
            text = pdfminer_extract_text(pdf_path)
            if text:
                # Write as markdown
                with open(md_path, 'w', encoding='utf-8') as f:
                    f.write(text)
                return True
            return False
            
        except Exception as e:
            self.log_message(f"Error in PDF to Markdown conversion: {e}")
            return False

    def convert_txt_to_markdown_file(self, txt_path, md_path):
        """Convert TXT to Markdown"""
        try:
            # Create output directory if needed
            output_dir = os.path.dirname(md_path)
            if output_dir:
                os.makedirs(output_dir, exist_ok=True)
                self.log_message(f"Created directory: {output_dir}")
            
            # Read text file
            with open(txt_path, 'r', encoding='utf-8', errors='ignore') as f:
                text = f.read()
            
            if text:
                # Write as markdown
                with open(md_path, 'w', encoding='utf-8') as f:
                    f.write(text)
                return True
            return False
            
        except Exception as e:
            self.log_message(f"Error in TXT to Markdown conversion: {e}")
            return False
            
    def log_message(self, message):
        """Add message to log"""
        timestamp = datetime.now().strftime("%H:%M:%S")
        log_entry = f"[{timestamp}] {message}\n"
        
        # Update log in main thread
        self.root.after(0, lambda: self.update_log(log_entry))
        
    def update_log(self, message):
        """Update the log text widget"""
        self.log_text.insert(tk.END, message)
        self.log_text.see(tk.END)
        
    def clear_log(self):
        """Clear the log"""
        self.log_text.delete(1.0, tk.END)
        
    def load_config(self):
        """Load configuration from file"""
        try:
            if os.path.exists(self.config_file):
                with open(self.config_file, 'r') as f:
                    config = json.load(f)
                    self.monitor_folders = config.get('monitor_folders', [])
                    self.docx_output_folder = config.get('docx_output_folder', '')
                    self.markdown_output_folder = config.get('markdown_output_folder', '')
                    self.monitor_interval.set(config.get('monitor_interval', 2))
        except Exception as e:
            print(f"Error loading config: {e}")
            
    def load_progress(self):
        """Load progress tracking from file"""
        try:
            if os.path.exists(self.progress_file):
                with open(self.progress_file, 'r') as f:
                    self.processed_files = json.load(f)
                    self.log_message(f"Loaded progress: {len(self.processed_files)} files tracked")
            else:
                self.processed_files = {}
        except Exception as e:
            print(f"Error loading progress: {e}")
            self.processed_files = {}
            
    def save_progress(self):
        """Save progress tracking to file"""
        try:
            with open(self.progress_file, 'w') as f:
                json.dump(self.processed_files, f, indent=2)
        except Exception as e:
            print(f"Error saving progress: {e}")
            
    def is_file_processed(self, file_path):
        """Check if a file has been successfully processed"""
        if file_path in self.processed_files:
            file_hash = self.get_file_hash(file_path)
            return self.processed_files[file_path].get('hash') == file_hash
        return False
        
    def mark_file_processed(self, file_path, status="completed"):
        """Mark a file as processed"""
        self.processed_files[file_path] = {
            'status': status,
            'hash': self.get_file_hash(file_path),
            'timestamp': datetime.now().isoformat()
        }
        self.save_progress()
        
    def mark_file_failed(self, file_path, error_msg=""):
        """Mark a file as failed"""
        self.processed_files[file_path] = {
            'status': 'failed',
            'hash': self.get_file_hash(file_path),
            'timestamp': datetime.now().isoformat(),
            'error': error_msg
        }
        self.save_progress()
        
    def clear_progress(self):
        """Clear all progress tracking"""
        self.processed_files = {}
        if os.path.exists(self.progress_file):
            os.remove(self.progress_file)
        self.log_message("🗑️ Progress tracking cleared - all files will be reprocessed")
        
    def get_progress_stats(self):
        """Get statistics about processed files"""
        total = len(self.processed_files)
        completed = sum(1 for f in self.processed_files.values() if f.get('status') == 'completed')
        failed = sum(1 for f in self.processed_files.values() if f.get('status') == 'failed')
        return total, completed, failed
        
    def show_progress_stats(self):
        """Show progress statistics in a popup"""
        total, completed, failed = self.get_progress_stats()
        
        stats_text = f"""📊 Progress Statistics

Total files tracked: {total}
✅ Successfully processed: {completed}
❌ Failed conversions: {failed}
⏭️ Skipped (already processed): {total - completed - failed}

Progress file: {self.progress_file}"""
        
        messagebox.showinfo("Progress Statistics", stats_text)
        
    def save_config(self):
        """Save configuration to file"""
        try:
            config = {
                'monitor_folders': self.monitor_folders,
                'docx_output_folder': self.docx_output_folder,
                'markdown_output_folder': self.markdown_output_folder,
                'monitor_interval': self.monitor_interval.get()
            }
            with open(self.config_file, 'w') as f:
                json.dump(config, f, indent=2)
        except Exception as e:
            print(f"Error saving config: {e}")
            
    def exit_app(self):
        """Exit the application"""
        if self.is_monitoring:
            self.stop_monitoring()
        self.save_config()
        self.save_progress()
        self.root.quit()
        
    def run(self):
        """Run the application"""
        self.root.mainloop()

def main():
    """Main function"""
    print("📁 Starting File Monitor & Auto-Converter...")
    app = FileMonitor()
    app.run()

if __name__ == "__main__":
    main() 