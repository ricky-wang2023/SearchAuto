#!/usr/bin/env python3
"""
Unified File Converter Tool
Combines DOC to DOCX and Markdown conversion features
"""

import os
import sys
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
import threading
from pathlib import Path
import re
import html
import subprocess
import shutil
from docx import Document
import PyPDF2
from pdfminer.high_level import extract_text as pdfminer_extract_text

class UnifiedConverter:
    def __init__(self):
        self.root = tk.Tk()
        self.root.title("🔄 Unified File Converter")
        self.root.geometry("900x700")
        self.root.configure(bg="#f5f5f5")
        self.root.resizable(True, True)
        
        # Variables
        self.input_folder = tk.StringVar()
        self.output_folder = tk.StringVar()
        self.conversion_progress = tk.DoubleVar()
        self.status_text = tk.StringVar(value="Ready to convert")
        
        # File type variables for markdown conversion
        self.docx_var = tk.BooleanVar(value=True)
        self.pdf_var = tk.BooleanVar(value=True)
        self.txt_var = tk.BooleanVar(value=True)
        
        self.setup_ui()
        
    def setup_ui(self):
        """Setup the user interface"""
        # Main frame
        main_frame = tk.Frame(self.root, bg="#f5f5f5")
        main_frame.pack(fill="both", expand=True, padx=20, pady=20)
        
        # Title
        title_label = tk.Label(main_frame, text="🔄 Unified File Converter", 
                             font=("Arial", 18, "bold"), bg="#f5f5f5", fg="#2196F3")
        title_label.pack(pady=(0, 20))
        
        # Create notebook for tabs
        notebook = ttk.Notebook(main_frame)
        notebook.pack(fill="both", expand=True, pady=10)
        
        # DOC to DOCX Tab
        doc_tab = tk.Frame(notebook, bg="#f5f5f5")
        notebook.add(doc_tab, text="📄 DOC to DOCX")
        self.setup_doc_tab(doc_tab)
        
        # Markdown Tab
        md_tab = tk.Frame(notebook, bg="#f5f5f5")
        notebook.add(md_tab, text="📝 Markdown Converter")
        self.setup_markdown_tab(md_tab)
        
        # Status frame (shared)
        status_frame = tk.LabelFrame(main_frame, text="🔄 Conversion Progress", 
                                   font=("Arial", 11, "bold"), bg="#f5f5f5", fg="navy")
        status_frame.pack(fill="x", pady=10)
        
        status_inner = tk.Frame(status_frame, bg="#f5f5f5")
        status_inner.pack(fill="x", padx=10, pady=10)
        
        self.progress_bar = ttk.Progressbar(status_inner, variable=self.conversion_progress, 
                                           maximum=100, length=500)
        self.progress_bar.pack(fill="x", pady=5)
        
        tk.Label(status_inner, textvariable=self.status_text, font=("Arial", 9), 
                bg="#f5f5f5").pack()
        
    def setup_doc_tab(self, parent):
        """Setup DOC to DOCX conversion tab"""
        # Input folder selection
        input_frame = tk.LabelFrame(parent, text="📁 Input Folder (DOC files)", 
                                  font=("Arial", 11, "bold"), bg="#f5f5f5", fg="navy")
        input_frame.pack(fill="x", pady=10, padx=10)
        
        input_inner = tk.Frame(input_frame, bg="#f5f5f5")
        input_inner.pack(fill="x", padx=10, pady=10)
        
        tk.Entry(input_inner, textvariable=self.input_folder, width=60, font=("Arial", 10)).pack(side="left", padx=(0, 10), fill="x", expand=True)
        tk.Button(input_inner, text="Browse", command=self.browse_input_folder, 
                 bg="#2196F3", fg="white", font=("Arial", 9, "bold")).pack(side="right")
        
        # Output folder selection
        output_frame = tk.LabelFrame(parent, text="📁 Output Folder (DOCX files)", 
                                   font=("Arial", 11, "bold"), bg="#f5f5f5", fg="navy")
        output_frame.pack(fill="x", pady=10, padx=10)
        
        output_inner = tk.Frame(output_frame, bg="#f5f5f5")
        output_inner.pack(fill="x", padx=10, pady=10)
        
        tk.Entry(output_inner, textvariable=self.output_folder, width=60, font=("Arial", 10)).pack(side="left", padx=(0, 10), fill="x", expand=True)
        tk.Button(output_inner, text="Browse", command=self.browse_output_folder, 
                 bg="#2196F3", fg="white", font=("Arial", 9, "bold")).pack(side="right")
        
        # Buttons frame
        buttons_frame = tk.Frame(parent, bg="#f5f5f5")
        buttons_frame.pack(fill="x", pady=20, padx=10)
        
        tk.Button(buttons_frame, text="🔄 Convert DOC to DOCX", command=self.start_doc_conversion, 
                 bg="#4CAF50", fg="white", font=("Arial", 12, "bold"), 
                 relief="flat", bd=0, height=2, width=20).pack(side="left", padx=5)
        
        tk.Button(buttons_frame, text="📊 Scan DOC Files", command=self.scan_doc_files, 
                 bg="#FF9800", fg="white", font=("Arial", 12, "bold"), 
                 relief="flat", bd=0, height=2, width=15).pack(side="left", padx=5)
        
    def setup_markdown_tab(self, parent):
        """Setup Markdown conversion tab"""
        # File types selection
        types_frame = tk.LabelFrame(parent, text="📄 File Types to Convert", 
                                  font=("Arial", 11, "bold"), bg="#f5f5f5", fg="navy")
        types_frame.pack(fill="x", pady=10, padx=10)
        
        types_inner = tk.Frame(types_frame, bg="#f5f5f5")
        types_inner.pack(fill="x", padx=10, pady=10)
        
        # Checkboxes for file types
        tk.Checkbutton(types_inner, text="DOCX files", variable=self.docx_var, 
                      bg="#f5f5f5", font=("Arial", 10)).pack(side="left", padx=10)
        tk.Checkbutton(types_inner, text="PDF files", variable=self.pdf_var, 
                      bg="#f5f5f5", font=("Arial", 10)).pack(side="left", padx=10)
        tk.Checkbutton(types_inner, text="TXT files", variable=self.txt_var, 
                      bg="#f5f5f5", font=("Arial", 10)).pack(side="left", padx=10)
        
        # Input folder selection
        input_frame = tk.LabelFrame(parent, text="📁 Input Folder", 
                                  font=("Arial", 11, "bold"), bg="#f5f5f5", fg="navy")
        input_frame.pack(fill="x", pady=10, padx=10)
        
        input_inner = tk.Frame(input_frame, bg="#f5f5f5")
        input_inner.pack(fill="x", padx=10, pady=10)
        
        tk.Entry(input_inner, textvariable=self.input_folder, width=60, font=("Arial", 10)).pack(side="left", padx=(0, 10), fill="x", expand=True)
        tk.Button(input_inner, text="Browse", command=self.browse_input_folder, 
                 bg="#2196F3", fg="white", font=("Arial", 9, "bold")).pack(side="right")
        
        # Output folder selection
        output_frame = tk.LabelFrame(parent, text="📁 Output Folder (Markdown files)", 
                                   font=("Arial", 11, "bold"), bg="#f5f5f5", fg="navy")
        output_frame.pack(fill="x", pady=10, padx=10)
        
        output_inner = tk.Frame(output_frame, bg="#f5f5f5")
        output_inner.pack(fill="x", padx=10, pady=10)
        
        tk.Entry(output_inner, textvariable=self.output_folder, width=60, font=("Arial", 10)).pack(side="left", padx=(0, 10), fill="x", expand=True)
        tk.Button(output_inner, text="Browse", command=self.browse_output_folder, 
                 bg="#2196F3", fg="white", font=("Arial", 9, "bold")).pack(side="right")
        
        # Buttons frame
        buttons_frame = tk.Frame(parent, bg="#f5f5f5")
        buttons_frame.pack(fill="x", pady=20, padx=10)
        
        tk.Button(buttons_frame, text="🔄 Convert to Markdown", command=self.start_markdown_conversion, 
                 bg="#9C27B0", fg="white", font=("Arial", 12, "bold"), 
                 relief="flat", bd=0, height=2, width=20).pack(side="left", padx=5)
        
        tk.Button(buttons_frame, text="📊 Scan Files", command=self.scan_markdown_files, 
                 bg="#FF9800", fg="white", font=("Arial", 12, "bold"), 
                 relief="flat", bd=0, height=2, width=15).pack(side="left", padx=5)
        
    def browse_input_folder(self):
        """Browse for input folder"""
        folder = filedialog.askdirectory(title="Select folder containing files to convert")
        if folder:
            self.input_folder.set(folder)
            
    def browse_output_folder(self):
        """Browse for output folder"""
        folder = filedialog.askdirectory(title="Select folder for converted files")
        if folder:
            self.output_folder.set(folder)
            
    def scan_doc_files(self):
        """Scan for DOC files in the input folder"""
        input_path = self.input_folder.get()
        if not input_path:
            messagebox.showwarning("Warning", "Please select an input folder first.")
            return
            
        if not os.path.exists(input_path):
            messagebox.showerror("Error", "Input folder does not exist.")
            return
            
        doc_files = []
        for root, dirs, files in os.walk(input_path):
            for file in files:
                if file.lower().endswith('.doc'):
                    doc_files.append(os.path.join(root, file))
                    
        if doc_files:
            messagebox.showinfo("Scan Results", f"Found {len(doc_files)} DOC files:\n\n" + 
                              "\n".join([os.path.basename(f) for f in doc_files[:10]]) + 
                              ("\n..." if len(doc_files) > 10 else ""))
        else:
            messagebox.showinfo("Scan Results", "No DOC files found in the selected folder.")
            
    def scan_markdown_files(self):
        """Scan for files to convert to markdown"""
        input_path = self.input_folder.get()
        if not input_path:
            messagebox.showwarning("Warning", "Please select an input folder first.")
            return
            
        if not os.path.exists(input_path):
            messagebox.showerror("Error", "Input folder does not exist.")
            return
            
        selected_types = self.get_selected_file_types()
        if not selected_types:
            messagebox.showwarning("Warning", "Please select at least one file type.")
            return
            
        files = []
        for root, dirs, filenames in os.walk(input_path):
            for file in filenames:
                if any(file.lower().endswith(ext) for ext in selected_types):
                    files.append(os.path.join(root, file))
                    
        if files:
            messagebox.showinfo("Scan Results", f"Found {len(files)} files to convert:\n\n" + 
                              "\n".join([os.path.basename(f) for f in files[:10]]) + 
                              ("\n..." if len(files) > 10 else ""))
        else:
            messagebox.showinfo("Scan Results", "No files found in the selected folder.")
            
    def get_selected_file_types(self):
        """Get list of selected file types"""
        types = []
        if self.docx_var.get():
            types.append('.docx')
        if self.pdf_var.get():
            types.append('.pdf')
        if self.txt_var.get():
            types.append('.txt')
        return types
        
    # DOC to DOCX conversion methods
    def convert_doc_to_docx(self, doc_path, docx_path):
        """Convert a single DOC file to DOCX"""
        try:
            # Normalize paths and handle encoding issues
            doc_path = os.path.abspath(doc_path)
            docx_path = os.path.abspath(docx_path)
            
            # Check if source file exists
            if not os.path.exists(doc_path):
                print(f"Source file not found: {doc_path}")
                return False
                
            # Create output directory if needed
            os.makedirs(os.path.dirname(docx_path), exist_ok=True)
            
            # Method 1: Try using LibreOffice (if available)
            if self.try_libreoffice_conversion(doc_path, docx_path):
                return True
                
            # Method 2: Try using Microsoft Word (if available)
            if self.try_word_conversion(doc_path, docx_path):
                return True
                
            # Method 3: Try using python-docx (limited success with old DOC files)
            if self.try_python_docx_conversion(doc_path, docx_path):
                return True
                
            # Method 4: Try to copy if it's already a DOCX in disguise
            if self.try_copy_if_docx(doc_path, docx_path):
                return True
                
            return False
            
        except Exception as e:
            print(f"Error converting {doc_path}: {e}")
            return False
            
    def try_libreoffice_conversion(self, doc_path, docx_path):
        """Try conversion using LibreOffice"""
        try:
            libreoffice_paths = [
                "soffice",  # Linux/Mac
                "libreoffice",  # Linux
                r"C:\Program Files\LibreOffice\program\soffice.exe",  # Windows
                r"C:\Program Files (x86)\LibreOffice\program\soffice.exe"  # Windows 32-bit
            ]
            
            for path in libreoffice_paths:
                try:
                    cmd = [path, "--headless", "--convert-to", "docx", 
                           "--outdir", os.path.dirname(docx_path), doc_path]
                    result = subprocess.run(cmd, capture_output=True, text=True, timeout=30)
                    
                    if result.returncode == 0:
                        expected_name = os.path.splitext(os.path.basename(doc_path))[0] + ".docx"
                        expected_path = os.path.join(os.path.dirname(docx_path), expected_name)
                        
                        if os.path.exists(expected_path):
                            if expected_path != docx_path:
                                shutil.move(expected_path, docx_path)
                            return True
                except:
                    continue
                    
            return False
            
        except Exception as e:
            print(f"LibreOffice conversion failed: {e}")
            return False
            
    def try_word_conversion(self, doc_path, docx_path):
        """Try conversion using Microsoft Word (Windows only)"""
        try:
            import win32com.client
            
            # Convert path to absolute path and fix slashes
            abs_doc_path = os.path.abspath(doc_path).replace('/', '\\')
            abs_docx_path = os.path.abspath(docx_path).replace('/', '\\')
            
            # Check if file exists
            if not os.path.exists(abs_doc_path):
                print(f"File not found: {abs_doc_path}")
                return False
            
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            
            try:
                doc = word.Documents.Open(abs_doc_path)
                doc.SaveAs2(abs_docx_path, FileFormat=16)  # 16 = docx format
                doc.Close()
                return True
            except Exception as e:
                print(f"Word document operation failed: {e}")
                return False
            finally:
                word.Quit()
            
        except Exception as e:
            print(f"Word conversion failed: {e}")
            return False
            
    def try_python_docx_conversion(self, doc_path, docx_path):
        """Try conversion using python-docx (limited success)"""
        try:
            from docx import Document
            
            # Check if file exists and is readable
            if not os.path.exists(doc_path):
                print(f"File not found: {doc_path}")
                return False
                
            # Try to read the file first
            try:
                with open(doc_path, 'rb') as f:
                    # Check if it's actually a DOCX file in disguise
                    content = f.read(4)
                    if content == b'PK\x03\x04':
                        # It's actually a DOCX file, just copy it
                        shutil.copy2(doc_path, docx_path)
                        return True
            except Exception as e:
                print(f"Cannot read file {doc_path}: {e}")
                return False
            
            # This method has limited success with old DOC files
            doc = Document(doc_path)
            doc.save(docx_path)
            return True
            
        except Exception as e:
            print(f"python-docx conversion failed: {e}")
            return False
            
    def try_copy_if_docx(self, doc_path, docx_path):
        """Check if file is actually a DOCX file in disguise and copy it"""
        try:
            with open(doc_path, 'rb') as f:
                # Check for ZIP file signature (DOCX files are ZIP archives)
                content = f.read(4)
                if content == b'PK\x03\x04':
                    print(f"File {os.path.basename(doc_path)} is actually a DOCX file, copying...")
                    shutil.copy2(doc_path, docx_path)
                    return True
            return False
        except Exception as e:
            print(f"Error checking file format: {e}")
            return False
            
    # Markdown conversion methods
    def convert_docx_to_markdown(self, docx_path, md_path):
        """Convert DOCX file to Markdown"""
        try:
            doc = Document(docx_path)
            markdown_content = []
            
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if not text:
                    markdown_content.append("")
                    continue
                    
                # Check for headings
                style_name = paragraph.style.name if paragraph.style and paragraph.style.name else ""
                if style_name.startswith('Heading'):
                    level_str = style_name[-1] if style_name[-1].isdigit() else "1"
                    level = int(level_str) if level_str.isdigit() else 1
                    markdown_content.append(f"{'#' * level} {text}")
                else:
                    # Check for bold and italic
                    formatted_text = text
                    for run in paragraph.runs:
                        if run.bold and run.italic:
                            formatted_text = formatted_text.replace(run.text, f"***{run.text}***")
                        elif run.bold:
                            formatted_text = formatted_text.replace(run.text, f"**{run.text}**")
                        elif run.italic:
                            formatted_text = formatted_text.replace(run.text, f"*{run.text}*")
                    
                    markdown_content.append(formatted_text)
                    
            # Write markdown file
            with open(md_path, 'w', encoding='utf-8') as f:
                f.write('\n'.join(markdown_content))
                
            return True
            
        except Exception as e:
            print(f"Error converting DOCX {docx_path}: {e}")
            return False
            
    def convert_pdf_to_markdown(self, pdf_path, md_path):
        """Convert PDF file to Markdown"""
        try:
            # Try pdfminer first (better for text extraction)
            try:
                text = pdfminer_extract_text(pdf_path)
            except:
                # Fallback to PyPDF2
                with open(pdf_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    text = ""
                    for page in pdf_reader.pages:
                        text += page.extract_text() + "\n"
                        
            # Clean and format text
            lines = text.split('\n')
            markdown_content = []
            
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                    
                # Simple heading detection (lines in caps or with numbers)
                if line.isupper() and len(line) < 100:
                    markdown_content.append(f"## {line}")
                elif re.match(r'^[0-9]+\.', line):
                    markdown_content.append(f"### {line}")
                else:
                    markdown_content.append(line)
                    
            # Write markdown file
            with open(md_path, 'w', encoding='utf-8') as f:
                f.write('\n'.join(markdown_content))
                
            return True
            
        except Exception as e:
            print(f"Error converting PDF {pdf_path}: {e}")
            return False
            
    def convert_txt_to_markdown(self, txt_path, md_path):
        """Convert TXT file to Markdown"""
        try:
            with open(txt_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
                
            # Clean and format text
            lines = content.split('\n')
            markdown_content = []
            
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                    
                # Simple heading detection
                if line.isupper() and len(line) < 100:
                    markdown_content.append(f"## {line}")
                elif re.match(r'^[0-9]+\.', line):
                    markdown_content.append(f"### {line}")
                else:
                    markdown_content.append(line)
                    
            # Write markdown file
            with open(md_path, 'w', encoding='utf-8') as f:
                f.write('\n'.join(markdown_content))
                
            return True
            
        except Exception as e:
            print(f"Error converting TXT {txt_path}: {e}")
            return False
            
    def convert_file_to_markdown(self, file_path, md_path):
        """Convert a single file to Markdown based on its type"""
        file_ext = os.path.splitext(file_path)[1].lower()
        
        if file_ext == '.docx':
            return self.convert_docx_to_markdown(file_path, md_path)
        elif file_ext == '.pdf':
            return self.convert_pdf_to_markdown(file_path, md_path)
        elif file_ext == '.txt':
            return self.convert_txt_to_markdown(file_path, md_path)
        else:
            print(f"Unsupported file type: {file_ext}")
            return False
            
    # Conversion start methods
    def start_doc_conversion(self):
        """Start DOC to DOCX conversion"""
        input_path = self.input_folder.get()
        output_path = self.output_folder.get()
        
        if not input_path or not output_path:
            messagebox.showwarning("Warning", "Please select both input and output folders.")
            return
            
        if not os.path.exists(input_path):
            messagebox.showerror("Error", "Input folder does not exist.")
            return
            
        if not os.path.exists(output_path):
            try:
                os.makedirs(output_path)
            except Exception as e:
                messagebox.showerror("Error", f"Cannot create output folder: {e}")
                return
                
        # Start conversion in a separate thread
        thread = threading.Thread(target=self.convert_doc_files)
        thread.daemon = True
        thread.start()
        
    def start_markdown_conversion(self):
        """Start Markdown conversion"""
        input_path = self.input_folder.get()
        output_path = self.output_folder.get()
        
        if not input_path or not output_path:
            messagebox.showwarning("Warning", "Please select both input and output folders.")
            return
            
        if not os.path.exists(input_path):
            messagebox.showerror("Error", "Input folder does not exist.")
            return
            
        if not os.path.exists(output_path):
            try:
                os.makedirs(output_path)
            except Exception as e:
                messagebox.showerror("Error", f"Cannot create output folder: {e}")
                return
                
        # Start conversion in a separate thread
        thread = threading.Thread(target=self.convert_markdown_files)
        thread.daemon = True
        thread.start()
        
    def convert_doc_files(self):
        """Convert all DOC files in the input folder"""
        input_path = self.input_folder.get()
        output_path = self.output_folder.get()
        
        # Find all DOC files
        doc_files = []
        for root, dirs, files in os.walk(input_path):
            for file in files:
                if file.lower().endswith('.doc'):
                    doc_files.append(os.path.join(root, file))
                    
        if not doc_files:
            self.root.after(0, lambda: messagebox.showinfo("Info", "No DOC files found to convert."))
            return
            
        # Update status
        self.root.after(0, lambda: self.status_text.set(f"Found {len(doc_files)} DOC files to convert..."))
        
        # Convert files
        successful = 0
        failed = 0
        
        for i, doc_file in enumerate(doc_files):
            # Update progress
            progress = (i / len(doc_files)) * 100
            self.root.after(0, lambda p=progress: self.conversion_progress.set(p))
            self.root.after(0, lambda f=os.path.basename(doc_file): 
                           self.status_text.set(f"Converting: {f}"))
            
            # Create output path
            rel_path = os.path.relpath(doc_file, input_path)
            docx_file = os.path.join(output_path, 
                                   os.path.splitext(rel_path)[0] + ".docx")
            
            # Create output directory if needed
            os.makedirs(os.path.dirname(docx_file), exist_ok=True)
            
            # Convert file
            if self.convert_doc_to_docx(doc_file, docx_file):
                successful += 1
            else:
                failed += 1
                
        # Update final status
        self.root.after(0, lambda: self.conversion_progress.set(100))
        self.root.after(0, lambda: self.status_text.set(
            f"DOC conversion complete! Success: {successful}, Failed: {failed}"))
        
        # Show results
        self.root.after(0, lambda: messagebox.showinfo("Conversion Complete", 
            f"DOC to DOCX conversion completed!\n\nSuccessful: {successful}\nFailed: {failed}\n\n"
            f"Output folder: {output_path}"))
            
    def convert_markdown_files(self):
        """Convert all files to Markdown"""
        input_path = self.input_folder.get()
        output_path = self.output_folder.get()
        selected_types = self.get_selected_file_types()
        
        # Find all files to convert
        files = []
        for root, dirs, filenames in os.walk(input_path):
            for file in filenames:
                if any(file.lower().endswith(ext) for ext in selected_types):
                    files.append(os.path.join(root, file))
                    
        if not files:
            self.root.after(0, lambda: messagebox.showinfo("Info", "No files found to convert."))
            return
            
        # Update status
        self.root.after(0, lambda: self.status_text.set(f"Found {len(files)} files to convert to Markdown..."))
        
        # Convert files
        successful = 0
        failed = 0
        
        for i, file_path in enumerate(files):
            # Update progress
            progress = (i / len(files)) * 100
            self.root.after(0, lambda p=progress: self.conversion_progress.set(p))
            self.root.after(0, lambda f=os.path.basename(file_path): 
                           self.status_text.set(f"Converting: {f}"))
            
            # Create output path
            rel_path = os.path.relpath(file_path, input_path)
            md_file = os.path.join(output_path, 
                                 os.path.splitext(rel_path)[0] + ".md")
            
            # Create output directory if needed
            os.makedirs(os.path.dirname(md_file), exist_ok=True)
            
            # Convert file
            if self.convert_file_to_markdown(file_path, md_file):
                successful += 1
            else:
                failed += 1
                
        # Update final status
        self.root.after(0, lambda: self.conversion_progress.set(100))
        self.root.after(0, lambda: self.status_text.set(
            f"Markdown conversion complete! Success: {successful}, Failed: {failed}"))
        
        # Show results
        self.root.after(0, lambda: messagebox.showinfo("Conversion Complete", 
            f"Markdown conversion completed!\n\nSuccessful: {successful}\nFailed: {failed}\n\n"
            f"Output folder: {output_path}"))
        
    def run(self):
        """Run the application"""
        self.root.mainloop()

def main():
    """Main function"""
    print("🔄 Starting Unified File Converter...")
    converter = UnifiedConverter()
    converter.run()

if __name__ == "__main__":
    main() 